"""Analyse uploaded template assets to drive the admin UI.

For a ``.docx`` we report the ``{{placeholders}}`` and ``{{#if}}`` conditionals
the renderer will act on, the paragraph styles the document actually defines
(so the style-mapping dropdowns can be populated), and which of the styles the
renderer relies on are missing. For an email ``.html`` we report the Mustache
variables and sections.

The reconciliation helper compares detected placeholders/conditionals against a
template's declared ``args`` so the UI can offer to add missing args and warn
about orphans.

This module is import-light and has no FastHTML dependency so it can be unit
tested on its own.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any, Dict, List

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from docx_tools.dynamic_docx_tools import PLACEHOLDER_PATTERN
from docx_tools.conditionals import parse_marker

# Styles the markdown renderer applies by name; warn when a template lacks them.
REQUIRED_DOCX_STYLES = [
    "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6",
    "List Bullet", "List Bullet 2", "List Bullet 3",
    "List Number", "List Number 2", "List Number 3",
    "Quote", "Table Grid", "Normal",
]

# Email standard fields are injected automatically by the email tool, so they
# never need a user-declared arg.
EMAIL_RESERVED_VARS = {"subject", "to", "cc", "bcc", "file_name"}

# Mustache section / variable patterns for HTML email templates.
_MUSTACHE_SECTION = re.compile(r"\{\{\s*[#^]\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
_MUSTACHE_VAR = re.compile(r"\{\{\{?\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}?\}\}")


@dataclass
class Analysis:
    """Result of analysing a template asset."""
    kind: str
    placeholders: List[str] = field(default_factory=list)
    conditionals: List[str] = field(default_factory=list)
    conditionals_balanced: bool = True
    styles_present: List[str] = field(default_factory=list)
    missing_required_styles: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


def _ordered_unique(items) -> List[str]:
    """De-duplicate *items* preserving first-seen order."""
    seen = set()
    out = []
    for it in items:
        if it not in seen:
            seen.add(it)
            out.append(it)
    return out


def _iter_docx_paragraph_texts(doc: DocxDocument):
    """Yield combined text of every paragraph in body, tables, headers, footers."""
    def _walk_container(paragraphs, tables):
        for p in paragraphs:
            yield p.text
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p.text

    yield from _walk_container(doc.paragraphs, doc.tables)
    for section in doc.sections:
        parts = [section.header, section.footer]
        if section.different_first_page_header_footer:
            parts += [section.first_page_header, section.first_page_footer]
        for part in parts:
            if part is None:
                continue
            yield from _walk_container(part.paragraphs, part.tables)


def _iter_docx_body_marker_texts(doc: DocxDocument):
    """Yield body-level paragraph texts (for conditional balance checking)."""
    p_tag = qn("w:p")
    t_tag = qn("w:t")
    for elem in doc.element.body:
        if elem.tag == p_tag:
            # Concatenate the paragraph's run text straight from the XML, avoiding
            # python-docx's internal Paragraph(elem, None) constructor.
            yield "".join(t.text or "" for t in elem.iter(t_tag))


def analyze_docx(data: bytes) -> Analysis:
    """Analyse a ``.docx`` given its bytes."""
    analysis = Analysis(kind="docx")
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as e:
        analysis.warnings.append(f"Could not open as a Word document: {e}")
        return analysis

    placeholders: List[str] = []
    for text in _iter_docx_paragraph_texts(doc):
        if "{{" in text:
            placeholders.extend(PLACEHOLDER_PATTERN.findall(text))
    analysis.placeholders = _ordered_unique(placeholders)

    # Conditionals + balance (body-level, matching the renderer's scope).
    conditionals: List[str] = []
    depth = 0
    balanced = True
    for text in _iter_docx_body_marker_texts(doc):
        marker = parse_marker(text)
        if marker is None:
            continue
        if marker.kind == "open":
            conditionals.append(marker.name)
            depth += 1
        else:
            depth -= 1
            if depth < 0:
                balanced = False
                depth = 0
    if depth != 0:
        balanced = False
    analysis.conditionals = _ordered_unique(conditionals)
    analysis.conditionals_balanced = balanced
    if not balanced:
        analysis.warnings.append(
            "Unbalanced {{#if}}/{{/if}} markers — every {{#if x}} needs a matching {{/if}}."
        )

    # Styles present vs required.
    try:
        present = {s.name for s in doc.styles if getattr(s, "name", None)}
    except Exception:
        present = set()
    analysis.styles_present = sorted(present)
    analysis.missing_required_styles = [s for s in REQUIRED_DOCX_STYLES if s not in present]

    return analysis


def analyze_html(data: bytes) -> Analysis:
    """Analyse an email ``.html`` template given its bytes."""
    analysis = Analysis(kind="email")
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("utf-8", errors="replace")
        analysis.warnings.append("File was not valid UTF-8; decoded with replacements.")

    sections = set(_MUSTACHE_SECTION.findall(text))
    analysis.conditionals = sorted(sections)

    variables = [
        v for v in _MUSTACHE_VAR.findall(text)
        if v not in sections and v not in EMAIL_RESERVED_VARS
    ]
    analysis.placeholders = _ordered_unique(variables)
    return analysis


def analyze(kind: str, data: bytes) -> Analysis:
    """Dispatch analysis by template *kind* (``docx`` or ``email``)."""
    if kind == "docx":
        return analyze_docx(data)
    if kind == "email":
        return analyze_html(data)
    raise ValueError(f"Unknown template kind: {kind!r}")


@dataclass
class Reconciliation:
    """Comparison of detected placeholders/conditionals against declared args."""
    missing_args: List[str] = field(default_factory=list)      # placeholder, no arg
    orphan_args: List[str] = field(default_factory=list)       # arg, no placeholder
    non_bool_conditions: List[str] = field(default_factory=list)  # condition arg not bool


def reconcile(analysis: Analysis, args: List[Dict[str, Any]]) -> Reconciliation:
    """Compare an :class:`Analysis` with a template's declared ``args``."""
    arg_by_name = {a.get("name"): a for a in (args or []) if isinstance(a, dict) and a.get("name")}
    arg_names = set(arg_by_name)

    detected = set(analysis.placeholders) | set(analysis.conditionals)
    rec = Reconciliation()
    rec.missing_args = [p for p in analysis.placeholders if p not in arg_names]
    # Conditionals also need a (bool) arg.
    rec.missing_args += [c for c in analysis.conditionals if c not in arg_names and c not in rec.missing_args]
    rec.orphan_args = [n for n in arg_by_name if n not in detected and n not in EMAIL_RESERVED_VARS]
    rec.non_bool_conditions = [
        c for c in analysis.conditionals
        if c in arg_by_name and str(arg_by_name[c].get("type", "string")).lower() not in ("bool", "boolean")
    ]
    return rec
