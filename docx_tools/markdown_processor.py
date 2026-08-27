"""Full markdown content processor (handles empty lines, soft breaks, blocks).
This module contains process_markdown_content and process_markdown_block which
orchestrate all block-level and inline parsing into a python-docx Document.
"""
import logging
from dataclasses import replace
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from .patterns import (
    HEADING_PATTERN,
    PAGE_BREAK_PATTERN,
    HORIZONTAL_LINE_PATTERN,
    IMAGE_PATTERN,
    TABLE_LINE_PATTERN,
    ORDERED_LIST_PATTERN,
    ORDERED_LIST_CAPTURE_PATTERN,
    UNORDERED_LIST_PATTERN,
    COMMENT_DIRECTIVE_PATTERN,
    CODE_FENCE_PATTERN,
    _ALIGN_CLOSE_RE,
    ordered_list_is_genuine,
    normalize_escaped_newlines,
    expand_br_to_block_breaks,
)
from .inline_formatting import parse_inline_formatting
from .block_elements import (
    parse_table,
    add_table_to_doc,
    process_list_items,
    add_horizontal_line,
    add_image_to_doc,
    detect_alignment,
)
from .style_map import (
    DEFAULT_STYLE_MAP,
    apply_style,
    add_mapped_heading,
    apply_style_to_block_element,
)
logger = logging.getLogger(__name__)


def _continues_ordered_run(stripped, ordered_run) -> bool:
    """True if the ordered marker on *stripped* continues the running count.

    *ordered_run* is a ``{'next': int | None}`` cell tracking the number that
    would continue the most recent top-level ordered list (see
    :func:`process_markdown_content`). This lets a continuation list resume after
    an intervening heading/blank line — e.g. items ``1.``/``2.`` under one
    heading and ``3.``/``4.`` under the next — even though, in isolation, a lone
    ``3.`` followed by a blank line is indistinguishable from a date and would
    not pass :func:`ordered_list_is_genuine`.
    """
    if not ordered_run or ordered_run.get('next') is None:
        return False
    match = ORDERED_LIST_CAPTURE_PATTERN.match(stripped)
    return bool(match) and int(match.group(1)) == ordered_run['next']


def process_markdown_content(doc, content, return_elements=False,
                             style_map=DEFAULT_STYLE_MAP):
    """Process full markdown content with all features: spacing, soft breaks, blocks.
    This is the single source of truth for converting a markdown string into
    document elements. Both the base tool and dynamic template placeholder
    replacement use this function.
    Args:
        doc: The python-docx Document instance.
        content: Raw markdown text (may contain newlines).
        return_elements: If True, created elements are detached from the doc body
            and returned (for reinsertion at a specific position).
    Returns:
        List of XML elements if return_elements is True, otherwise an empty list.
    """
    # Treat literal "\n"/"\r\n" sequences (typed as text rather than real
    # newlines) as genuine newlines so they split into paragraphs/blocks instead
    # of being mangled into stray "n" characters downstream.
    content = normalize_escaped_newlines(content)
    # Promote <br> that borders block content (lists/headings) to real newlines
    # so such blocks are detected; a prose <br> stays an inline soft break.
    content = expand_br_to_block_breaks(content)
    lines = content.split('\n')
    n = len(lines)
    i = 0
    all_elements = []
    # Running ordered-list count: the number that would continue the most recent
    # top-level ordered list. Preserved across headings, blank lines, block
    # quotes and comment-directive blocks so a list can resume after a section
    # heading or an interposed quote/styled note; reset by any other block
    # content. Lets _continues_ordered_run() accept e.g. "3." after a heading
    # even when it is blank-separated (and so not locally genuine). A mutable
    # cell so process_list_items can update it through process_markdown_block.
    ordered_run = {'next': None}
    while i < n:
        line = lines[i]
        # --- Empty line handling (preserve spacing) ---
        if not line.strip():
            empty_line_count = 1
            i += 1
            while i < n and not lines[i].strip():
                empty_line_count += 1
                i += 1
            if empty_line_count >= 2:
                for _ in range(empty_line_count - 1):
                    p = doc.add_paragraph()
                    if return_elements:
                        all_elements.append(p._p)
                        doc._body._body.remove(p._p)
            continue
        # --- Soft line breaks (trailing two spaces) ---
        # A complete <!-- --> comment line is exempt: trailing spaces after the
        # closing marker must not turn a directive (or a comment to be skipped)
        # into soft-break prose that renders the comment text literally.
        if line.endswith('  ') and not _is_complete_comment(line):
            paragraph_lines = []
            while i < n:
                current_line = lines[i]
                if not current_line.strip():
                    break
                paragraph_lines.append(current_line)
                i += 1
                if not current_line.endswith('  '):
                    break
            full_text = '  \n'.join(paragraph_lines)
            first_line = paragraph_lines[0].strip()
            if first_line.startswith('#'):
                stripped_hashes = first_line.lstrip('#')
                level = len(first_line) - len(stripped_hashes)
                elem = _add_heading(doc, level, stripped_hashes.strip(), style_map)._p
                # A heading does not break ordered-list continuation.
            elif first_line.startswith('>'):
                elem = _add_quote(doc, full_text[1:].strip(), style_map)._p
                # A quote does not break ordered-list continuation (like a
                # heading, it deliberately interrupts a numbered run).
            else:
                para = doc.add_paragraph()
                parse_inline_formatting(full_text, para)
                elem = para._p
                ordered_run['next'] = None
            if return_elements:
                all_elements.append(elem)
                doc._body._body.remove(elem)
            continue
        # --- All other block elements: delegate to block processor ---
        # Continuation survives headings, blank lines (above), block quotes and
        # complete <!-- --> comment lines. A style directive attaches to the
        # block it styles, and that WHOLE directive-styled block (whatever its
        # type — the dispatcher consumes directive + target together) is treated
        # as a deliberate interruption of a numbered run, exactly like a
        # heading: e.g. an evidence note or citation between numbered
        # paragraphs of a legal filing. Any other block content breaks the run.
        # A numbered line is left alone so a list that actually renders can
        # update the count via process_list_items. An unclosed "<!--" is not a
        # comment (it renders as literal prose), so it resets like any prose.
        stripped = line.strip()
        is_comment_line = _is_complete_comment(line)
        if (HEADING_PATTERN.match(stripped) is None
                and ORDERED_LIST_PATTERN.match(stripped) is None
                and not stripped.startswith('>')
                and not is_comment_line):
            ordered_run['next'] = None
        i, block_elems = process_markdown_block(doc, lines, i,
                                                return_element=return_elements,
                                                style_map=style_map,
                                                ordered_run=ordered_run)
        if return_elements:
            all_elements.extend(block_elems)
    return all_elements
_CODE_FONT = 'Courier New'


def _add_heading(doc, level, content, style_map):
    """Create a heading paragraph (mapped style) and parse *content* into it.

    Shared by the block dispatcher and the soft-break path so heading rendering
    lives in one place.
    """
    heading = add_mapped_heading(doc, min(level, 6), style_map)
    parse_inline_formatting(content, heading)
    return heading


def _is_complete_comment(line):
    """True if *line* (ignoring surrounding whitespace) is a whole ``<!-- -->``."""
    stripped = line.strip()
    return stripped.startswith('<!--') and stripped.endswith('-->')


def _has_explicit_style(element):
    """True if *element* is a paragraph carrying an explicit ``w:pStyle``.

    Used by the style-directive branch to spot list items that were already
    styled through the overridden style map. Note python-docx drops ``w:pStyle``
    when a style application falls back to the document default ("Normal"), so
    such paragraphs are (harmlessly) re-styled by the caller.
    """
    if element.tag != qn('w:p'):
        return False
    ppr = element.find(qn('w:pPr'))
    return ppr is not None and ppr.find(qn('w:pStyle')) is not None


def _add_quote(doc, content, style_map):
    """Create a block-quote paragraph (mapped style) and parse *content* into it."""
    para = doc.add_paragraph()
    apply_style(para, style_map.quote)
    parse_inline_formatting(content, para)
    return para


def _render_code_block(doc, lines, start_idx, fence_match, style_map, collect):
    """Render a fenced code block verbatim as monospace paragraphs.

    *fence_match* is the opener match. Consumes lines up to and including the
    closing fence (a line of the same fence character, at least as long, with no
    info string). Each code line becomes one paragraph so blank lines and
    indentation are preserved; markdown inside is intentionally NOT parsed.
    Returns the index of the first line after the block.
    """
    fence = fence_match.group(1)
    fence_char = fence[0]
    fence_len = len(fence)
    n = len(lines)
    j = start_idx + 1
    while j < n:
        closing = lines[j].strip()
        if closing and set(closing) == {fence_char} and len(closing) >= fence_len:
            j += 1  # consume the closing fence
            break
        para = doc.add_paragraph()
        # add_run preserves leading/trailing whitespace via xml:space="preserve"
        run = para.add_run(lines[j])
        if style_map.code:
            # Use the mapped paragraph style's font; a run-level override would
            # otherwise always win over the style's monospace font.
            apply_style(para, style_map.code, fallback=None)
            if para.style.name != style_map.code:
                # Mapped style is missing from the template — keep it monospace.
                run.font.name = _CODE_FONT
        else:
            run.font.name = _CODE_FONT
        collect(para._p)
        j += 1
    return j


def process_markdown_block(doc, lines, start_idx, return_element=True,
                           style_map=DEFAULT_STYLE_MAP, directives=None,
                           ordered_run=None):
    """Process a single markdown block element and return created XML elements.

    *directives* carries comment-directive options (`borderless`, `widths`, …)
    collected from `<!-- … -->` lines immediately above this block; see the
    directive branch below.

    *ordered_run* is the running ordered-list count cell from
    :func:`process_markdown_content`; when present it lets a numbered line that
    continues the previous list (e.g. after a heading) start a list even if it is
    not locally genuine, and lets the rendered list update the count.
    Returns:
        Tuple of (next_index, list_of_elements).
    """
    line = lines[start_idx]
    stripped = line.strip()
    elements = []
    def _collect(element):
        """If return_element, detach *element* from body and collect it."""
        if return_element:
            elements.append(element)
            doc._body._body.remove(element)
    try:
        # Heading
        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading = _add_heading(doc, level, heading_match.group(2), style_map)
            _collect(heading._p)
            return start_idx + 1, elements
        # Fenced code block (``` or ~~~) — content is taken verbatim, NOT parsed
        # as markdown, so headings/lists/backticks inside code are preserved.
        fence_match = CODE_FENCE_PATTERN.match(stripped)
        if fence_match:
            next_idx = _render_code_block(doc, lines, start_idx, fence_match,
                                          style_map, _collect)
            return next_idx, elements
        # Table (lines starting with |)
        if TABLE_LINE_PATTERN.match(stripped):
            table_data, col_alignments, next_idx = parse_table(lines, start_idx)
            if table_data:
                # Table options come from comment directives collected above the
                # table (see the directive branch below).
                d = directives or {}
                borderless = 'borderless' in d
                col_widths = None
                if 'widths' in d:
                    try:
                        col_widths = [float(v) for v in d['widths'].split()]
                    except ValueError:
                        col_widths = None
                word_table = add_table_to_doc(table_data, doc,
                                             col_alignments=col_alignments,
                                             borderless=borderless,
                                             col_widths=col_widths,
                                             table_style=style_map.table)
                if word_table is not None:
                    _collect(word_table._tbl)
                return next_idx, elements
        # Page break (---)
        if PAGE_BREAK_PATTERN.match(stripped):
            doc.add_page_break()
            _collect(doc.paragraphs[-1]._p)
            return start_idx + 1, elements
        # Horizontal line (***)
        if HORIZONTAL_LINE_PATTERN.match(stripped):
            _collect(add_horizontal_line(doc)._p)
            return start_idx + 1, elements
        # Image (![alt](url))
        img_match = IMAGE_PATTERN.match(stripped)
        if img_match:
            body = doc._body._body
            existing_children = list(body) if return_element else None
            add_image_to_doc(doc, img_match.group(2), img_match.group(1))
            if return_element:
                for element in list(body)[len(existing_children):]:
                    elements.append(element)
                    body.remove(element)
            return start_idx + 1, elements
        # Alignment (inline or block-open)
        align_result = detect_alignment(stripped)
        if align_result is not None:
            inner, alignment = align_result
            if inner is not None:
                # Inline form: <center>text</center> / <div align="...">text</div>.
                # A heading written inside it still becomes a real heading (then
                # aligned), rather than rendering its "#" marks as literal text.
                # Only headings are promoted here; a single-line list such as
                # "<center>- item</center>" stays inline prose — use the multi-line
                # block form (handled by _process_alignment_block) for lists.
                heading_match = HEADING_PATTERN.match(inner)
                if heading_match:
                    para = _add_heading(doc, len(heading_match.group(1)),
                                        heading_match.group(2), style_map)
                else:
                    para = doc.add_paragraph()
                    parse_inline_formatting(inner, para)
                para.alignment = alignment
                _collect(para._p)
                return start_idx + 1, elements
            # Block-open form: render the inner lines through the full block
            # pipeline so headings/lists inside the block are recognised, then
            # stamp the alignment on every produced paragraph.
            idx, block_elems = _process_alignment_block(
                doc, lines, start_idx + 1, alignment, style_map,
                return_element, ordered_run,
            )
            if return_element and block_elems:
                elements.extend(block_elems)
            return idx, elements
        # Ordered list. A numbered line only starts a list when it begins at 1 or
        # has a continuation (see ordered_list_is_genuine); otherwise it falls
        # through to a plain paragraph so a standalone date like "23. června 2026"
        # is not misread as an ordered list.
        if ORDERED_LIST_PATTERN.match(stripped) and (
                ordered_list_is_genuine(lines, start_idx)
                or _continues_ordered_run(stripped, ordered_run)):
            return process_list_items(
                lines, start_idx, doc, is_ordered=True, level=0, return_elements=return_element,
                number_styles=style_map.list_number, bullet_styles=style_map.list_bullet,
                ordered_run=ordered_run,
            )
        # Unordered list
        if UNORDERED_LIST_PATTERN.match(stripped):
            return process_list_items(
                lines, start_idx, doc, is_ordered=False, level=0, return_elements=return_element,
                number_styles=style_map.list_number, bullet_styles=style_map.list_bullet,
            )
        # Blockquote (> text)
        if stripped.startswith('>'):
            quote_para = _add_quote(doc, stripped[1:].strip(), style_map)
            _collect(quote_para._p)
            return start_idx + 1, elements
        # Comment directives: <!-- borderless -->, <!-- widths: … -->, <!-- style: … -->.
        # Collect consecutive directive lines and attach them to the next block
        # (single look-ahead mechanism for all block directives).
        directive_match = COMMENT_DIRECTIVE_PATTERN.match(stripped)
        if directive_match:
            collected = dict(directives) if directives else {}
            idx = start_idx
            while idx < len(lines):
                m = COMMENT_DIRECTIVE_PATTERN.match(lines[idx].strip())
                if not m:
                    break
                collected[m.group(1).lower()] = (m.group(2) or '').strip()
                idx += 1
            # Skip blank lines between the directives and the block they modify.
            while idx < len(lines) and not lines[idx].strip():
                idx += 1
            if idx >= len(lines):
                return idx, elements  # directives with nothing to attach to → no-op
            body = doc._body._body
            # Snapshot existing children as a set (holding the proxies alive so lxml
            # keeps stable identities, and giving O(1) membership). New elements are
            # inserted before the trailing <w:sectPr>, so positional slicing would be
            # unreliable.
            existing = None if return_element else set(body)
            # A style directive targeting a LIST is applied through the style map
            # (top level only) rather than stamped onto every produced paragraph
            # afterwards: nested items keep their own mapped level styles instead
            # of being flattened, and ordered-list numbering resolves from the
            # directive style itself, keeping its numeral format and indents.
            style_name = collected.get('style')
            target = lines[idx].strip()
            target_is_ordered = bool(ORDERED_LIST_PATTERN.match(target))
            styles_via_map = bool(style_name) and (
                target_is_ordered or bool(UNORDERED_LIST_PATTERN.match(target)))
            block_style_map = style_map
            if styles_via_map:
                if target_is_ordered:
                    block_style_map = replace(
                        style_map,
                        list_number=(style_name,) + tuple(style_map.list_number[1:]))
                else:
                    block_style_map = replace(
                        style_map,
                        list_bullet=(style_name,) + tuple(style_map.list_bullet[1:]))
            new_idx, block_elems = process_markdown_block(
                doc, lines, idx, return_element=return_element,
                style_map=block_style_map, directives=collected,
                ordered_run=ordered_run,
            )
            # The 'style' directive applies the named style to whatever was produced.
            if style_name:
                produced = (block_elems if return_element
                            else [el for el in body if el not in existing])
                for el in produced:
                    # A list item already styled via the overridden style map (or a
                    # nested level's own mapped style) carries an explicit pStyle —
                    # leave it alone. A numbered line that rendered as plain prose
                    # (e.g. a standalone date) has none and still gets the style.
                    if styles_via_map and _has_explicit_style(el):
                        continue
                    apply_style_to_block_element(doc, el, style_name)
            if return_element:
                elements.extend(block_elems)
            return new_idx, elements
        # Other HTML comments (not a recognised directive) — skip silently.
        if stripped.startswith('<!--') and stripped.endswith('-->'):
            return start_idx + 1, elements
        # Regular paragraph
        para = doc.add_paragraph()
        parse_inline_formatting(stripped, para)
        _collect(para._p)
        return start_idx + 1, elements
    except Exception as e:
        logger.error("Failed to process markdown block at line %d: %s", start_idx, e, exc_info=True)
        return start_idx + 1, elements


def _process_alignment_block(doc, lines, start_idx, alignment, style_map,
                             return_element, ordered_run):
    """Render the lines inside a multi-line ``<center>``/``<div align>`` block.

    Each inner line goes through the normal block pipeline (so headings, lists,
    tables, etc. are recognised instead of rendering their markers as literal
    text), and *alignment* is then applied to every produced paragraph.

    Returns ``(next_index, produced_elements)``. ``produced_elements`` is
    populated only when *return_element* is True (the block pipeline detaches the
    elements for the caller); otherwise the produced paragraphs stay in the body
    and are located via a before/after snapshot so the alignment can be stamped.
    """
    body = doc._body._body
    existing = None if return_element else set(body)
    collected = []
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if _ALIGN_CLOSE_RE.match(stripped):
            i += 1
            break
        if not stripped:
            i += 1
            continue
        i, produced = process_markdown_block(
            doc, lines, i, return_element=return_element,
            style_map=style_map, ordered_run=ordered_run,
        )
        if return_element:
            collected.extend(produced)
    produced_all = (collected if return_element
                    else [el for el in body if el not in existing])
    for el in produced_all:
        if el.tag == qn('w:p'):  # alignment applies to paragraphs, not tables
            Paragraph(el, doc._body).alignment = alignment
    return i, collected
