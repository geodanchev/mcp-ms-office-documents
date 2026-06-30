"""Dynamic registration of DOCX template MCP tools.

Similar to dynamic email templates, this module allows defining custom DOCX templates
with placeholders ({{placeholder}}) and YAML configuration for template-specific arguments.

Placeholders in DOCX templates use Mustache syntax:
  - {{placeholder}} - replaced with markdown-formatted text
  - Text supports inline markdown: **bold**, *italic*, `code`, [links](url)

YAML configuration example:
```yaml
templates:
  - name: formal_letter
    description: Formal business letter template
    docx_path: letter_template.docx  # filename only, searched in custom/default templates
    annotations:
      title: Formal Letter Generator
    args:
      - name: recipient_name
        type: string
        description: Full name of the recipient
        required: true
      - name: body
        type: string
        description: Main body text (supports markdown formatting)
        required: true
```
"""
from __future__ import annotations

import io
import re
import copy
import logging
import threading
from pathlib import Path
from typing import Any, Dict, Optional, Literal

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from pydantic import Field, create_model
from fastmcp import FastMCP

from upload_tools import upload_file
from template_utils import find_file_in_template_dirs
from template_registry import gather_specs, safe_remove_tool
from async_runner import run_blocking
import metrics
from .conditionals import resolve_conditionals
from .inline_formatting import parse_inline_formatting
from .patterns import (
    contains_block_markdown, normalize_escaped_newlines, expand_br_to_block_breaks,
)
from .markdown_processor import process_markdown_content
from .style_map import DEFAULT_STYLE_MAP, build_style_map
from fastmcp.exceptions import ToolError


__all__ = [
    "register_docx_template_tools_from_yaml",
    "register_docx_template",
    "unregister_docx_template",
    "registered_docx_template_names",
    "replace_placeholders_in_document",
]

logger = logging.getLogger(__name__)

# Live registry of docx template tools registered on the running server,
# mapping tool name -> the spec it was built from. Lets the admin UI list,
# replace and remove dynamic docx tools after startup. Guarded by a lock because
# it is mutated from HTTP request handlers (admin save/delete) and read from the
# status page concurrently.
_REGISTERED_DOCX: Dict[str, Dict[str, Any]] = {}
_REG_LOCK = threading.Lock()

# Type mapping for YAML -> Python types
TYPE_MAP = {
    "string": str, "str": str,
    "int": int, "integer": int,
    "float": float,
    "bool": bool, "boolean": bool,
    "list": list[str], "list[str]": list[str], "list[string]": list[str],
}

# Regex to find Mustache-style placeholders: {{name}} or {{{name}}}
PLACEHOLDER_PATTERN = re.compile(r'\{\{\{?([a-zA-Z_][a-zA-Z0-9_]*)\}?\}\}')



def _insert_markdown_content_after_paragraph(
    doc: DocxDocument,
    paragraph: Paragraph,
    content: str,
    style_map=DEFAULT_STYLE_MAP,
) -> list:
    """Insert markdown content (including lists, headings, soft breaks) after a paragraph.

    Uses the unified process_markdown_content() which handles all markdown
    features: empty-line spacing, soft line breaks (trailing two spaces),
    and all block-level elements.

    Args:
        doc: The Word document
        paragraph: The paragraph after which to insert content
        content: The markdown content to insert
        style_map: Style mapping for rendered block content (see StyleMap).

    Returns:
        The list of inserted XML elements (``<w:p>`` / ``<w:tbl>``), in order, so
        the caller can post-process them (e.g. propagate the placeholder format).
    """
    try:
        # Process content and get detached elements
        new_elements = process_markdown_content(doc, content, return_elements=True,
                                                style_map=style_map)

        # Find the paragraph's position in the document body and insert after it
        body = doc._body._body
        p_element = paragraph._p
        para_idx = list(body).index(p_element)

        for offset, elem in enumerate(new_elements, start=1):
            body.insert(para_idx + offset, elem)
        return new_elements
    except Exception as e:
        logger.error("Failed to insert markdown content after paragraph: %s", e, exc_info=True)
        return []


def find_docx_template_by_name(filename: str) -> Optional[str]:
    """Find a specific DOCX template by filename in custom/default template directories.

    Args:
        filename: The filename of the DOCX template (e.g., 'letter_template.docx')

    Returns:
        Absolute path to the template file as string, or None if not found.
    """
    found = find_file_in_template_dirs(filename)
    return str(found) if found else None


def _copy_run_format(src_run, dst_run) -> None:
    """Copy *src_run*'s character formatting onto *dst_run*.

    Deep-copies the source run properties element (``<w:rPr>``) so every direct
    format — bold, italic, underline, colour, highlight, font, character style —
    is preserved wholesale, rather than enumerating individual properties.
    """
    src_rpr = src_run._r.find(qn('w:rPr'))
    if src_rpr is None:
        return
    dst_rpr = dst_run._r.find(qn('w:rPr'))
    if dst_rpr is not None:
        dst_run._r.remove(dst_rpr)
    dst_run._r.insert(0, copy.deepcopy(src_rpr))  # rPr must be the run's first child


def _add_formatted_segments(paragraph, segments) -> None:
    """Append *segments* (``(text, source_run)`` pairs) as runs, keeping format."""
    for seg_text, seg_run in segments:
        _copy_run_format(seg_run, paragraph.add_run(seg_text))


def _propagate_format_to_block(doc, inserted, src_ppr, fmt) -> None:
    """Give *inserted* default-styled paragraphs the placeholder's look.

    Paragraphs the markdown styled deliberately (headings, list items, quotes,
    code — i.e. those with an explicit ``<w:pStyle>``) are left untouched. Plain
    (default/Normal) paragraphs inherit the placeholder paragraph's *layout*
    (alignment/indent/spacing) and the placeholder run's character format *fmt*,
    so e.g. a justified letter body stays justified instead of falling back to the
    document default.
    """
    # Layout-only copy of the placeholder's paragraph properties: drop its named
    # style (<w:pStyle>) and list numbering (<w:numPr>) so neither is stamped onto
    # produced prose (e.g. a placeholder that sits in a Heading or a list item).
    # Produced prose keeps its own default style and inherits only direct layout
    # (alignment/indent/spacing).
    layout_ppr = None
    if src_ppr is not None:
        layout_ppr = copy.deepcopy(src_ppr)
        for tag in ('w:pStyle', 'w:numPr'):
            child = layout_ppr.find(qn(tag))
            if child is not None:
                layout_ppr.remove(child)
    for elem in inserted:
        if elem.tag != qn('w:p'):
            continue  # tables etc. are not styled here
        para = Paragraph(elem, doc._body)
        ppr = elem.find(qn('w:pPr'))
        # An explicit paragraph style means the markdown set the look on purpose.
        if ppr is not None and ppr.find(qn('w:pStyle')) is not None:
            continue
        if layout_ppr is not None:
            # Preserve a deliberate markdown alignment (e.g. <center>, <div align>)
            # across the pPr swap so it wins over the placeholder's own alignment;
            # re-apply via the API so <w:jc> lands in schema order.
            own_alignment = para.alignment
            if ppr is not None:
                elem.remove(ppr)
            elem.insert(0, copy.deepcopy(layout_ppr))  # pPr must be the paragraph's first child
            if own_alignment is not None:
                para.alignment = own_alignment
        for run in para.runs:
            _apply_placeholder_format(run, fmt)


def _segments_for_range(run_info, lo: int, hi: int):
    """Return ``(text, run)`` pairs for the part of each run within ``[lo, hi)``.

    Splits the surrounding text at the placeholder boundary while remembering
    which original run each slice came from, so its formatting can be restored.
    """
    segments = []
    for start, end, run in run_info:
        seg_lo = max(start, lo)
        seg_hi = min(end, hi)
        if seg_lo < seg_hi:
            segments.append((run.text[seg_lo - start:seg_hi - start], run))
    return segments


def _apply_placeholder_format(run, fmt) -> None:
    """Fill *run*'s unset formatting from the placeholder's captured format *fmt*.

    Only properties the markdown value left unset (``None``) are filled, so
    formatting the value asked for (e.g. ``**bold**``) is never clobbered.
    """
    if fmt['name'] and not run.font.name:
        run.font.name = fmt['name']
    if fmt['size'] and not run.font.size:
        run.font.size = fmt['size']
    if fmt['color_rgb'] and not run.font.color.rgb:
        run.font.color.rgb = fmt['color_rgb']
    elif fmt['color_theme'] and not run.font.color.theme_color:
        run.font.color.theme_color = fmt['color_theme']
    # Emphasis formats are tri-state (None = inherit). Propagate the placeholder's
    # explicit value — including an explicit False (e.g. bold turned off to
    # counteract a paragraph style) — but never override what the markdown set.
    if fmt['bold'] is not None and run.bold is None:
        run.bold = fmt['bold']
    if fmt['italic'] is not None and run.italic is None:
        run.italic = fmt['italic']
    if fmt['underline'] is not None and run.underline is None:
        run.underline = fmt['underline']
    if fmt['highlight'] and run.font.highlight_color is None:
        run.font.highlight_color = fmt['highlight']


def _replace_placeholder_in_paragraph(
    paragraph: Paragraph,
    placeholder: str,
    value: str,
    doc: DocxDocument = None,
    style_map=DEFAULT_STYLE_MAP,
) -> bool:
    """Replace a placeholder in a paragraph with markdown-formatted text.

    This function handles the case where a placeholder might be split across multiple runs
    (which Word often does when editing documents).

    For block-level content (lists), the content is inserted as new paragraphs after
    the current paragraph.

    Args:
        paragraph: The paragraph to search and modify
        placeholder: The placeholder text including braces (e.g., '{{name}}')
        value: The replacement value (supports markdown formatting)
        doc: The Word document (required for block-level content like lists)

    Returns:
        True if replacement was made, False otherwise
    """
    try:
        # First, try to find the placeholder in the full paragraph text
        full_text = paragraph.text
        if placeholder not in full_text:
            return False

        # Normalise literal "\n"/"\r\n" the model may have written as text (not
        # real newlines) BEFORE routing, so block detection and the multi-line
        # check below see genuine line breaks. process_markdown_content /
        # parse_inline_formatting normalise again downstream (idempotent — that is
        # what the base tool relies on), so this early pass is for routing only.
        value = normalize_escaped_newlines(value)
        # Likewise promote a <br> that borders block content (a list/heading) to a
        # real newline up front, so contains_block_markdown / the multi-line check
        # route such values to the block pipeline instead of inline-only (where the
        # list would render as literal text). Prose <br> is left as a soft break.
        value = expand_br_to_block_breaks(value)

        # Collect all runs and their text
        runs = list(paragraph.runs)
        if not runs:
            return False

        # Build a map of character positions to runs
        combined_text = ""
        run_info = []  # List of (start_pos, end_pos, run)

        for run in runs:
            start = len(combined_text)
            combined_text += run.text
            end = len(combined_text)
            run_info.append((start, end, run))

        # Find the placeholder in the combined text
        placeholder_start = combined_text.find(placeholder)
        if placeholder_start == -1:
            return False

        placeholder_end = placeholder_start + len(placeholder)

        # Store formatting from the run where placeholder starts
        formatting_run = None
        for start, end, run in run_info:
            if start <= placeholder_start < end:
                formatting_run = run
                break

        # Capture the placeholder run's direct character formatting so it can be
        # re-applied to the replacement text (font, colour, and the emphasis
        # formats bold/italic/underline/highlight).
        placeholder_format = {
            'name': formatting_run.font.name if formatting_run else None,
            'size': formatting_run.font.size if formatting_run else None,
            'color_rgb': formatting_run.font.color.rgb if formatting_run else None,
            'color_theme': formatting_run.font.color.theme_color if formatting_run else None,
            'bold': formatting_run.bold if formatting_run else None,
            'italic': formatting_run.italic if formatting_run else None,
            'underline': formatting_run.underline if formatting_run else None,
            'highlight': formatting_run.font.highlight_color if formatting_run else None,
        }

        # Strategy: Rebuild the paragraph content
        # 1. Get text before placeholder
        # 2. Get replacement content (parsed markdown)
        # 3. Get text after placeholder

        # Slice the surrounding text at the placeholder boundaries, remembering
        # each slice's source run so its formatting can be restored (rather than
        # flattening before/after text to plain runs).
        before_segments = _segments_for_range(run_info, 0, placeholder_start)
        after_segments = _segments_for_range(run_info, placeholder_end, len(combined_text))

        # Check if the value contains block-level content (lists, headings)
        has_block_content = contains_block_markdown(value)

        # A MULTI-LINE whole-paragraph placeholder (no surrounding text) is rendered
        # through the SAME markdown pipeline as the base tool, so newline handling is
        # identical everywhere: a line break starts a new paragraph; soft breaks use
        # <br> or two trailing spaces. Block content is always routed there too. A
        # single-line value, a placeholder in the MIDDLE of a paragraph, or one in a
        # table/header (doc is None) stays inline (soft breaks only) — a single line
        # needs no splitting and a sentence cannot be split into separate paragraphs.
        is_whole_paragraph = not before_segments and not after_segments
        use_block = (
            doc is not None
            and value.strip() != ""
            and (has_block_content or (is_whole_paragraph and "\n" in value))
        )

        # Capture the placeholder paragraph's own properties so produced paragraphs
        # can inherit them (see _propagate_format_to_block).
        existing_ppr = paragraph._p.find(qn('w:pPr'))
        src_ppr = copy.deepcopy(existing_ppr) if existing_ppr is not None else None

        # Clear all existing runs
        p_element = paragraph._p
        for run in runs:
            p_element.remove(run._r)

        # Re-add the text before the placeholder, preserving its formatting.
        _add_formatted_segments(paragraph, before_segments)

        if use_block:
            # Render via the base-tool pipeline and give the produced plain
            # paragraphs the placeholder's style/format (headings/lists keep theirs).
            inserted = _insert_markdown_content_after_paragraph(doc, paragraph, value, style_map)
            _propagate_format_to_block(doc, inserted or [], src_ppr, placeholder_format)

            # Re-add the text after the placeholder, preserving its formatting.
            _add_formatted_segments(paragraph, after_segments)

            # If the placeholder occupied the whole paragraph it is now empty –
            # remove it so the produced paragraphs take its place cleanly.
            if is_whole_paragraph:
                p_element.getparent().remove(p_element)
        else:
            # Inline replacement: parse the value into runs of THIS paragraph
            # (newlines / <br> become soft breaks), then fill any formatting the
            # markdown left unset from the placeholder run's captured format.
            runs_before_value = len(paragraph.runs)
            parse_inline_formatting(value, paragraph)
            for run in list(paragraph.runs)[runs_before_value:]:
                _apply_placeholder_format(run, placeholder_format)

            # Re-add the text after the placeholder, preserving its formatting.
            _add_formatted_segments(paragraph, after_segments)

        return True

    except Exception as e:
        logger.error("Failed to replace placeholder '%s' in paragraph: %s", placeholder, e, exc_info=True)
        return False


def _replace_placeholders_in_paragraph(
    paragraph: Paragraph,
    context: Dict[str, str],
    doc: DocxDocument = None,
    style_map=DEFAULT_STYLE_MAP,
) -> None:
    """Replace all placeholders in a paragraph with their values.

    This function iteratively replaces placeholders one at a time, re-scanning
    the paragraph after each replacement to handle position shifts correctly.

    Args:
        paragraph: The paragraph to process
        context: Dictionary mapping placeholder names to their values
        doc: The Word document (required for block-level content like lists)
    """
    # Keep replacing until no more placeholders are found
    max_iterations = 100  # Safety limit to prevent infinite loops
    iteration = 0

    while iteration < max_iterations:
        iteration += 1

        # Get current paragraph text and find placeholders
        full_text = paragraph.text
        matches = PLACEHOLDER_PATTERN.findall(full_text)

        if not matches:
            break

        # Find the first placeholder that exists in context
        replaced = False
        for placeholder_name in matches:
            if placeholder_name not in context:
                continue

            value = context[placeholder_name]
            if value is None:
                value = ""
            else:
                value = str(value)

            # Try triple brace first, then double brace
            for placeholder in [f'{{{{{{{placeholder_name}}}}}}}', f'{{{{{placeholder_name}}}}}']:
                if placeholder in paragraph.text:
                    if _replace_placeholder_in_paragraph(paragraph, placeholder, value, doc,
                                                         style_map):
                        replaced = True
                        break

            if replaced:
                break  # Re-scan paragraph after successful replacement

        if not replaced:
            # No more replaceable placeholders found
            break


def _replace_placeholders_in_table(
    table: Table,
    context: Dict[str, str],
    doc: DocxDocument = None,
    style_map=DEFAULT_STYLE_MAP,
) -> None:
    """Replace all placeholders in a table.

    Note: Block-level content (lists) is not supported in table cells.

    Args:
        table: The table to process
        context: Dictionary mapping placeholder names to their values
        doc: The Word document (not used for tables, as block content not supported)
        style_map: Style mapping for any rendered inline content.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Note: We don't pass doc to avoid inserting lists in table cells
                _replace_placeholders_in_paragraph(paragraph, context, doc=None,
                                                   style_map=style_map)


def _replace_placeholders_in_document(doc: DocxDocument, context: Dict[str, str],
                                      style_map=DEFAULT_STYLE_MAP) -> None:
    """Replace all placeholders in the entire document.

    Processes:
    - Main document body paragraphs
    - Tables in the main body
    - Headers and footers

    Args:
        doc: The Word document to process
        context: Dictionary mapping placeholder names to their values
    """
    # Process main body paragraphs
    for paragraph in doc.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, context, doc, style_map)

    # Process tables
    for table in doc.tables:
        _replace_placeholders_in_table(table, context, doc, style_map)

    # Process headers and footers
    for section in doc.sections:
        # Collect all header/footer parts to process
        parts = []

        # Default header and footer
        if section.header:
            parts.append(section.header)
        if section.footer:
            parts.append(section.footer)

        # First-page header/footer (when template uses "Different First Page")
        if section.different_first_page_header_footer:
            if section.first_page_header:
                parts.append(section.first_page_header)
            if section.first_page_footer:
                parts.append(section.first_page_footer)

        # Even-page header/footer (when template uses "Different Even & Odd Pages")
        even_page_header = getattr(section, 'even_page_header', None)
        even_page_footer = getattr(section, 'even_page_footer', None)
        if even_page_header:
            parts.append(even_page_header)
        if even_page_footer:
            parts.append(even_page_footer)

        for part in parts:
            for paragraph in part.paragraphs:
                # Headers/footers: don't support block content
                _replace_placeholders_in_paragraph(paragraph, context, doc=None,
                                                   style_map=style_map)
            for table in part.tables:
                _replace_placeholders_in_table(table, context, doc=None,
                                               style_map=style_map)


# Public alias: the admin preview renderer reuses the document substitution
# pipeline across the package boundary.
replace_placeholders_in_document = _replace_placeholders_in_document


def docx_spec_dir(yaml_path: Path) -> Path:
    """Directory holding UI-managed per-template docx specs, beside *yaml_path*."""
    return yaml_path.parent / "docx_templates.d"


def registered_docx_template_names() -> list[str]:
    """Return the names of docx template tools currently registered."""
    with _REG_LOCK:
        return sorted(_REGISTERED_DOCX)


def register_docx_template_tools_from_yaml(mcp: FastMCP, yaml_path: Path) -> None:
    """Register dynamic DOCX template tools from YAML.

    Merges the master YAML at *yaml_path* with any UI-managed per-template files
    in the sibling ``docx_templates.d/`` directory (the per-template file wins on
    a name clash). Either source may be absent.

    Args:
        mcp: The FastMCP instance to register tools with
        yaml_path: Path to the master YAML configuration file
    """
    templates, cfg = gather_specs(yaml_path, docx_spec_dir(yaml_path))
    global_style_mapping = cfg.get("style_mapping") or {}

    for spec in templates:
        try:
            register_docx_template(mcp, spec, global_style_mapping)
        except Exception as e:
            name = spec.get("name", "<unknown>") if isinstance(spec, dict) else "<unknown>"
            logger.exception(f"[dynamic-docx] Failed to register template '{name}': {e}")


def register_docx_template(
    mcp: FastMCP, spec: Dict[str, Any], global_style_mapping: Dict[str, Any] = None
) -> bool:
    """Register (or replace) a single docx template tool on the live server.

    If a tool with the same name already exists it is removed first, so this is
    safe to call after startup to apply an edited template. Returns True on
    success.
    """
    name = spec.get("name") if isinstance(spec, dict) else None
    if name:
        safe_remove_tool(mcp, name)
        with _REG_LOCK:
            _REGISTERED_DOCX.pop(name, None)
    return _register_single_template(mcp, spec, global_style_mapping)


def unregister_docx_template(mcp: FastMCP, name: str) -> bool:
    """Remove a docx template tool from the live server. Returns True if removed."""
    removed = safe_remove_tool(mcp, name)
    with _REG_LOCK:
        _REGISTERED_DOCX.pop(name, None)
    return removed


def _register_single_template(mcp: FastMCP, spec: Dict[str, Any],
                              global_style_mapping: Dict[str, Any] = None) -> bool:
    """Register a single DOCX template as an MCP tool.

    Args:
        mcp: The FastMCP instance
        spec: The template specification from YAML
        global_style_mapping: Document-wide ``style_mapping`` overrides; the
            template's own ``style_mapping`` (if any) takes precedence.

    Returns:
        True when the tool was registered; False when the spec was skipped
        (missing name/path, file not found, …).
    """
    name = spec.get("name")
    if not name:
        logger.warning("[dynamic-docx] Template missing 'name', skipping.")
        return False

    description = spec.get("description", f"Generate document from {name} template")
    annotations = spec.get("annotations", {})
    docx_path = spec.get("docx_path")

    if not docx_path:
        logger.warning(f"[dynamic-docx] Missing docx_path for {name}, skipping.")
        return False

    # Validate path is filename only (no directory components)
    docx_path_obj = Path(docx_path)
    if docx_path_obj.is_absolute() or len(docx_path_obj.parts) != 1:
        logger.error(
            f"[dynamic-docx] docx_path must be filename only (no directories) for {name}; "
            f"got '{docx_path}'"
        )
        return False

    # Resolve the template file
    resolved = find_docx_template_by_name(docx_path)
    if not resolved:
        logger.error(f"[dynamic-docx] Template file not found for {name}: {docx_path}")
        return False

    logger.info(f"[dynamic-docx] Using template for {name}: {resolved}")

    # Resolve the style map for this template (per-template overrides global).
    style_map = build_style_map(global_style_mapping, spec.get("style_mapping"))

    # Build Pydantic model fields from args
    fields: Dict[str, Any] = {}

    for arg in spec.get("args", []):
        arg_name = arg.get("name")
        if not arg_name:
            continue

        # Handle enum values
        enum_values = arg.get("enum")
        if enum_values and isinstance(enum_values, list) and enum_values:
            if all(isinstance(v, int) for v in enum_values):
                lit_values = tuple(int(v) for v in enum_values)
            elif all(isinstance(v, (int, float)) for v in enum_values):
                lit_values = tuple(float(v) for v in enum_values)
            else:
                lit_values = tuple(str(v) for v in enum_values)
            py_type = Literal[lit_values]  # type: ignore[index]
            required = bool(arg.get("required", True))
            default = arg.get("default", (... if required else None))
            if default is not ... and default is not None and default not in lit_values:
                logger.warning(
                    f"[dynamic-docx] Default '{default}' not in enum for {arg_name}; ignoring default."
                )
                default = ... if required else None
            desc = arg.get("description") or f"One of: {', '.join(map(str, lit_values))}"
            fields[arg_name] = (py_type, Field(default, description=desc))
            continue

        # Handle regular types
        py_type = TYPE_MAP.get(str(arg.get("type", "string")).lower(), str)
        required = bool(arg.get("required", True))
        field_type = py_type if required else Optional[py_type]  # type: ignore[index]
        default = arg.get("default", (... if required else None))
        desc = arg.get("description", "")
        fields[arg_name] = (field_type, Field(default, description=desc) if desc else default)

    # Create the Pydantic model
    model = create_model(f"{name}_DocxArgs", **fields)  # type: ignore
    # Expose the dynamically-created model in the module globals: Pydantic/FastMCP
    # resolve the tool's annotation (`data: <Model>`) by name against this module's
    # namespace when building the tool schema, so the class must be importable from
    # here. The name `<tool>_DocxArgs` is unique per template; re-registering an
    # edited template intentionally overwrites the prior model.
    globals()[model.__name__] = model

    # Create the tool function.
    #
    # The tool body (`_sync_impl`) is synchronous and performs blocking
    # work: opening the .docx zip, mustache-style placeholder
    # substitution, and synchronous upload to the configured backend.
    # It is wrapped in an `async def` (`tool_impl`) that dispatches the
    # call through `run_blocking()`, so the work either runs on a
    # worker thread (when RUN_BLOCKING_BY_ASYNCIO_THREAD_ENABLED is
    # truthy) or inline on the event loop (default, legacy behaviour).
    # FastMCP awaits the async tool directly, leaving dispatch entirely
    # to our helper — keeping behaviour consistent with the static tools
    # in main.py.
    def make_tool_fn(_model=model, _template_path=resolved, _name=name,
                     _style_map=style_map):
        def _sync_impl(data):
            try:
                # Load the template document
                doc = DocxDocument(_template_path)

                # Build context from input data
                payload = data.model_dump()

                # Resolve conditional blocks ({{#if flag}} ... {{/if}}) before
                # substitution, since this prunes whole block elements.
                resolve_conditionals(doc, payload)

                context = {k: ("" if v is None else str(v)) for k, v in payload.items()}

                # Replace placeholders
                _replace_placeholders_in_document(doc, context, _style_map)

                # Save to buffer and upload
                buffer = io.BytesIO()
                try:
                    doc.save(buffer)
                    buffer.seek(0)

                    result = upload_file(buffer, "docx", filename=payload.get("file_name") or _name)
                finally:
                    buffer.close()

                logger.info(f"[dynamic-docx] Document generated from template {_name}")
                metrics.record_call("docx", _name)
                return result

            except Exception as e:
                metrics.record_error("docx", _name, str(e))
                logger.error(f"[dynamic-docx] Error generating document from {_name}: {e}", exc_info=True)
                raise ToolError(f"Error generating document from template {_name}: {e}")

        async def tool_impl(data: _model) -> str:  # type: ignore
            return await run_blocking(_sync_impl, data)

        tool_impl.__annotations__['data'] = _model  # type: ignore[index]
        tool_impl.__annotations__['return'] = str  # type: ignore[index]
        return tool_impl

    # Register the tool
    mcp.tool(
        name=name,
        description=description,
        annotations=annotations,
        tags={"docx", "document", "template"},
    )(make_tool_fn())

    with _REG_LOCK:
        _REGISTERED_DOCX[name] = spec
    logger.info(f"[dynamic-docx] Registered tool: {name}")
    return True

