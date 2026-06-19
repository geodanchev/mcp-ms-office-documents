"""Block-level markdown elements: tables, lists, images, alignment, horizontal lines."""
import re
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .patterns import (
    ORDERED_LIST_PATTERN,
    UNORDERED_LIST_PATTERN,
    ORDERED_LIST_CAPTURE_PATTERN,
    UNORDERED_LIST_CAPTURE_PATTERN,
    _ALIGN_INLINE_RE,
    _ALIGN_OPEN_RE,
    _ALIGN_CLOSE_RE,
)
from .inline_formatting import parse_inline_formatting
from .patterns import _BR_RE
from .numbering import (
    resolve_ordered_abstract_num_id,
    new_restarted_num,
    apply_numbering,
)
from .style_map import apply_style

logger = logging.getLogger(__name__)


ALIGNMENT_MAP = {
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
}
# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------
_SEPARATOR_RE = re.compile(r'^[|:\-\s]+$')

def _parse_alignment_row(line):
    """Extract column alignments from a markdown table separator row.
    Returns a list of alignment values (WD_ALIGN_PARAGRAPH) or None per column.
    """
    cells = [c.strip() for c in line.split('|')[1:-1]]
    alignments = []
    for cell in cells:
        cell = cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif cell.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            alignments.append(None)  # left is default, no need to set explicitly
    return alignments

def parse_table(lines, start_idx):
    """Parse markdown table and return table data, column alignments, and next line index.
    Returns:
        Tuple of (table_data, col_alignments, next_line_index).
        table_data is a list of rows (each row is a list of cell strings).
        col_alignments is a list of WD_ALIGN_PARAGRAPH values (or None) per column.
    """
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
        else:
            break
    if len(table_lines) < 2:
        return None, None, start_idx + 1
    table_data = []
    col_alignments = None
    for line in table_lines:
        # Detect separator row and extract alignment
        if _SEPARATOR_RE.match(line.replace('|', ' | ')):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if all(re.match(r'^:?-+:?$', c.strip()) for c in cells if c.strip()):
                col_alignments = _parse_alignment_row(line)
                continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        table_data.append(cells)
    return table_data, col_alignments, i

def _remove_table_borders(table):
    """Remove all borders from a Word table (makes it invisible)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    # Remove existing borders element if present
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def add_table_to_doc(table_data, doc, col_alignments=None, borderless=False,
                     col_widths=None, table_style='Table Grid'):
    """Add table data to Word document.
    Args:
        table_data: List of rows, each a list of cell text strings.
        doc: python-docx Document instance.
        col_alignments: Optional list of WD_ALIGN_PARAGRAPH per column.
        borderless: If True, remove all table borders (invisible table).
        col_widths: Optional list of proportional width values per column.
            Values are normalized to sum to 100% of available page width.
        table_style: Word table style name to apply (falls back to the document
            default if the named style is missing).
    Returns the created ``Table`` object, or ``None`` when the table could
    not be created (empty data or exception).
    """
    if not table_data:
        return None
    rows = len(table_data)
    cols = max(len(row) for row in table_data) if table_data else 0
    try:
        word_table = doc.add_table(rows=rows, cols=cols)
    except Exception as e2:
        logger.error("Failed to create table: %s", e2, exc_info=True)
        return None
    apply_style(word_table, table_style, fallback=None)
    if borderless:
        _remove_table_borders(word_table)
    # Apply column widths if specified
    if col_widths and len(col_widths) >= cols:
        total = sum(col_widths[:cols])
        page_width = 6.5  # usable width in inches (standard letter with 1" margins)
        for j in range(cols):
            width_inches = (col_widths[j] / total) * page_width
            for row in word_table.rows:
                row.cells[j].width = Inches(width_inches)
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                try:
                    cell = word_table.cell(i, j)
                    # Split on <br> variants to create multiple paragraphs in cell
                    segments = _BR_RE.split(cell_text)
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                    parse_inline_formatting(segments[0], cell.paragraphs[0])
                    for seg in segments[1:]:
                        new_para = cell.add_paragraph()
                        parse_inline_formatting(seg.strip(), new_para)
                    # Apply column alignment to all paragraphs in cell
                    if col_alignments and j < len(col_alignments) and col_alignments[j]:
                        for para in cell.paragraphs:
                            para.alignment = col_alignments[j]
                except Exception as e:
                    logger.warning("Failed to populate table cell [%d, %d]: %s", i, j, e)
    return word_table
# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------
def process_list_items(lines, start_idx, doc, is_ordered=False, level=0,
                       return_elements=False, number_styles=None, bullet_styles=None,
                       base_indent=None):
    """Process markdown list items with proper Word numbering.
    When *return_elements* is True the created paragraph XML elements are
    removed from the document body and returned so the caller can re-insert
    them elsewhere (used by the template placeholder machinery).
    *number_styles*/*bullet_styles* override the per-level Word style names
    (see :class:`docx_tools.style_map.StyleMap`); they default to the built-ins.

    Nesting is determined by *relative* indentation: items sharing the first
    item's indent (*base_indent*) are siblings at *level*; a more-indented item
    begins a child list one level deeper. This makes any consistent indent unit
    work (2, 3 or 4 spaces, or a tab) rather than assuming a fixed step.

    Ordered-list numbering restarts only when the explicit number ``1`` reappears
    at a level — NOT on any backward jump. ``1, 2, 1`` restarts; ``3, 4, 3`` does
    not (Word renders ``3, 4, 5``). A list starting at *N* honours that as the
    first value via ``startOverride``.
    Returns:
        Tuple of (next_line_index, list_of_elements | None).
    """
    if bullet_styles is None:
        bullet_styles = ('List Bullet', 'List Bullet 2', 'List Bullet 3')
    if number_styles is None:
        number_styles = ('List Number', 'List Number 2', 'List Number 3')
    style_array = number_styles if is_ordered else bullet_styles
    style = style_array[min(level, len(style_array) - 1)]
    elements = [] if return_elements else None
    i = start_idx
    n = len(lines)
    list_capture_pattern = (
        ORDERED_LIST_CAPTURE_PATTERN if is_ordered else UNORDERED_LIST_CAPTURE_PATTERN
    )
    # Ordered-list numbering restart state (see docx_tools/numbering.py). Each logical
    # list gets its own <w:num> instance so it counts independently; a new instance is
    # started for the first item and whenever "1." reappears at this level.
    abstract_num_id = None
    numbering_root = None
    num_resolved = False
    current_num_id = None
    items_emitted = 0
    while i < n:
        original_line = lines[i]
        stripped_left = original_line.lstrip()
        indent = len(original_line) - len(stripped_left)
        line = stripped_left.rstrip()
        if base_indent is None:
            base_indent = indent  # first item defines this level's indent
        # Items more/less indented than this level are handled by the caller
        # (a child list via the look-ahead below, or an ancestor on dedent).
        if indent != base_indent:
            break
        list_match = list_capture_pattern.match(line)
        if not list_match:
            break
        if is_ordered:
            item_number = int(list_match.group(1))
            item_text = list_match.group(2)
        else:
            item_number = None
            item_text = list_match.group(1)
        paragraph = doc.add_paragraph()
        apply_style(paragraph, style, fallback='Normal')
        if is_ordered:
            if current_num_id is None or (item_number == 1 and items_emitted > 0):
                if not num_resolved:
                    abstract_num_id, numbering_root = resolve_ordered_abstract_num_id(doc)
                    num_resolved = True
                if abstract_num_id is not None:
                    current_num_id = new_restarted_num(
                        numbering_root, abstract_num_id, level, start=item_number
                    )
            if current_num_id is not None:
                apply_numbering(paragraph, current_num_id, level)
        parse_inline_formatting(item_text, paragraph)
        items_emitted += 1
        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)
        i += 1
        # Look ahead for nested items
        while i < n:
            next_original = lines[i]
            next_stripped_left = next_original.lstrip()
            next_line = next_stripped_left.rstrip()
            if not next_line:
                i += 1
                continue
            next_indent = len(next_original) - len(next_stripped_left)
            if next_indent > base_indent:
                is_nested_ordered = bool(ORDERED_LIST_PATTERN.match(next_line))
                is_nested_unordered = bool(UNORDERED_LIST_PATTERN.match(next_line))
                if is_nested_ordered or is_nested_unordered:
                    i, nested = process_list_items(
                        lines, i, doc, is_nested_ordered, level + 1, return_elements,
                        number_styles=number_styles, bullet_styles=bullet_styles,
                        base_indent=next_indent,
                    )
                    if return_elements and nested:
                        elements.extend(nested)
                else:
                    break
            else:
                break
    # Forward-progress guarantee
    if i == start_idx:
        original_line = lines[start_idx]
        stripped_line = original_line.strip()
        logger.warning(
            "process_list_items: no progress at line %d (%r) for level=%d; "
            "rendering as a plain paragraph to guarantee forward progress",
            start_idx, stripped_line, level,
        )
        paragraph = doc.add_paragraph()
        parse_inline_formatting(stripped_line, paragraph)
        if return_elements:
            elements.append(paragraph._p)
            doc._body._body.remove(paragraph._p)
        i = start_idx + 1
    return i, elements
# ---------------------------------------------------------------------------
# Page break / horizontal line
# ---------------------------------------------------------------------------
def add_horizontal_line(doc):
    """Add a visual horizontal line (thin border) to the document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p
# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------
def add_image_to_doc(doc, url, alt_text, max_width_inches=None):
    """Add an image from a URL to the document.
    Downloads the image and inserts it.  On failure inserts an error
    placeholder paragraph instead.
    """
    try:
        from pptx_tools.image_utils import download_image
        if max_width_inches is None:
            try:
                sec = doc.sections[-1]
                max_width_inches = (sec.page_width - sec.left_margin - sec.right_margin) / 914400
            except Exception:
                max_width_inches = 5.5
        image_stream, _ = download_image(url)
        doc.add_picture(image_stream, width=Inches(max_width_inches))
        if alt_text:
            caption = doc.add_paragraph()
            caption.add_run(alt_text).italic = True
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.warning("Failed to add image from '%s': %s", url, e)
        doc.add_paragraph().add_run(f"[Image could not be loaded: {url}]")
# ---------------------------------------------------------------------------
# Text alignment
# ---------------------------------------------------------------------------
def detect_alignment(line):
    """Detect an alignment tag (inline *or* block-open) on *line*.
    Returns ``(inner_text, alignment)`` for an inline tag,
    ``(None, alignment)`` for a block-open tag, or ``None`` if no match.
    """
    m = _ALIGN_INLINE_RE.match(line)
    if m:
        if m.group(1) is not None:
            return m.group(1).strip(), WD_ALIGN_PARAGRAPH.CENTER
        return m.group(3).strip(), ALIGNMENT_MAP.get(m.group(2).lower(), WD_ALIGN_PARAGRAPH.LEFT)
    m = _ALIGN_OPEN_RE.match(line)
    if m:
        align = ALIGNMENT_MAP.get((m.group(1) or 'center').lower(), WD_ALIGN_PARAGRAPH.CENTER)
        return None, align
    return None


def process_alignment_block(lines, start_idx, doc, alignment, return_elements=False):
    """Process lines inside a multi-line alignment block."""
    elements = [] if return_elements else None
    i = start_idx
    while i < len(lines):
        stripped = lines[i].strip()
        if _ALIGN_CLOSE_RE.match(stripped):
            i += 1
            break
        if not stripped:
            i += 1
            continue
        para = doc.add_paragraph()
        para.alignment = alignment
        parse_inline_formatting(stripped, para)
        if return_elements:
            elements.append(para._p)
            doc._body._body.remove(para._p)
        i += 1
    return i, elements
