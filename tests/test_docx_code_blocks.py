"""Tests for fenced code blocks (improvement 4.1).

A ` ``` ` / ` ~~~ ` fenced block is rendered verbatim as monospace paragraphs;
markdown inside it (headings, lists, backticks) is intentionally NOT parsed.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.markdown_processor import process_markdown_content  # noqa: E402
from docx_tools.style_map import build_style_map  # noqa: E402
from docx_tools.patterns import contains_block_markdown  # noqa: E402


def _new(start_styles=()):
    d = Document()
    for s in start_styles:
        d.styles.add_style(s, WD_STYLE_TYPE.PARAGRAPH)
    return d


def _lines(doc, start):
    return [p for p in doc.paragraphs[start:]]


def test_code_lines_are_monospace_and_not_parsed():
    doc = _new()
    start = len(doc.paragraphs)
    md = (
        "```python\n"
        "# not a heading\n"
        "- not a bullet\n"
        "1. not a list\n"
        "```\n"
    )
    process_markdown_content(doc, md)
    rendered = [(p.text, p.style.name) for p in _lines(doc, start) if p.text]
    # Verbatim text, all Normal paragraphs (no Heading/List styles).
    assert rendered == [
        ("# not a heading", "Normal"),
        ("- not a bullet", "Normal"),
        ("1. not a list", "Normal"),
    ]
    for p in _lines(doc, start):
        if p.text:
            assert all(r.font.name == "Courier New" for r in p.runs)


def test_leading_whitespace_and_blank_lines_preserved():
    doc = _new()
    start = len(doc.paragraphs)
    md = "```\ndef f():\n    return 1\n\n    return 2\n```\n"
    process_markdown_content(doc, md)
    texts = [p.text for p in _lines(doc, start)]
    assert "    return 1" in texts          # 4-space indent kept
    assert "" in texts                       # blank line inside block preserved
    assert "    return 2" in texts


def test_tilde_fence_and_unclosed_at_eof():
    doc = _new()
    start = len(doc.paragraphs)
    # No closing fence — should consume to end of input (CommonMark behaviour).
    process_markdown_content(doc, "~~~\nline one\nline two\n")
    rendered = [(p.text, [r.font.name for r in p.runs]) for p in _lines(doc, start) if p.text]
    assert rendered == [("line one", ["Courier New"]), ("line two", ["Courier New"])]


def test_backticks_inside_code_are_literal():
    doc = _new()
    start = len(doc.paragraphs)
    process_markdown_content(doc, "```\nuse `code` spans here\n```\n")
    rendered = [p.text for p in _lines(doc, start) if p.text]
    assert rendered == ["use `code` spans here"]  # backticks NOT consumed as inline code


def test_code_style_is_mappable():
    doc = _new(start_styles=("My Code",))
    start = len(doc.paragraphs)
    sm = build_style_map({"code": "My Code"})
    process_markdown_content(doc, "```\nx = 1\n```\n", style_map=sm)
    styled = [p for p in _lines(doc, start) if p.text]
    assert styled and all(p.style.name == "My Code" for p in styled)


def test_mapped_code_style_font_not_overridden():
    """A mapped code style keeps its own font (no run-level Courier New override)."""
    doc = _new(start_styles=("Mono Code",))
    start = len(doc.paragraphs)
    process_markdown_content(doc, "```\nx = 1\n```\n", style_map=build_style_map({"code": "Mono Code"}))
    p = [p for p in _lines(doc, start) if p.text == "x = 1"][0]
    assert p.style.name == "Mono Code"
    # No run-level font is forced, so the style's font wins.
    assert all(r.font.name is None for r in p.runs)


def test_default_code_block_uses_courier_new():
    """With no mapped code style, runs fall back to the monospace font."""
    doc = _new()
    start = len(doc.paragraphs)
    process_markdown_content(doc, "```\ny = 2\n```\n")
    p = [p for p in _lines(doc, start) if p.text == "y = 2"][0]
    assert all(r.font.name == "Courier New" for r in p.runs)


def test_missing_mapped_code_style_falls_back_to_monospace():
    """If the mapped code style is absent from the template, stay monospace."""
    doc = _new()  # "Code Block" style intentionally NOT registered
    start = len(doc.paragraphs)
    process_markdown_content(doc, "```\nz = 3\n```\n", style_map=build_style_map({"code": "Code Block"}))
    p = [p for p in _lines(doc, start) if p.text == "z = 3"][0]
    assert p.style.name != "Code Block"          # style was missing
    assert all(r.font.name == "Courier New" for r in p.runs)  # but still monospace


def test_code_fence_detected_as_block_markdown():
    # Ensures dynamic-template placeholder values with code render as block content.
    assert contains_block_markdown("```\ncode\n```")


def test_content_around_code_block_still_renders():
    doc = _new()
    start = len(doc.paragraphs)
    process_markdown_content(doc, "# Title\n\n```\ncode\n```\n\nAfter.\n")
    pairs = [(p.text, p.style.name) for p in _lines(doc, start) if p.text]
    assert ("Title", "Heading 1") in pairs
    assert ("code", "Normal") in pairs
    assert ("After.", "Normal") in pairs
