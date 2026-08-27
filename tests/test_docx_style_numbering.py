"""Tests for style-aware ordered-list numbering and indents.

Three behaviours introduced together (follow-up to issues #66/#67):

1. Restart numbering instances are resolved from the ordered-list *style
   actually applied* at each level — a ``style_mapping`` override or a
   ``<!-- style: … -->`` directive — so restarted lists keep that style's own
   numeral format; nested levels use the level style's own declared
   ``numId``/``ilvl`` instead of assuming a multi-level ``List Number`` abstract.
2. Any ``w:ind`` the applied style defines is re-asserted as direct paragraph
   formatting, because a direct ``w:numPr`` otherwise lets the numbering level's
   indents silently override the style's.
3. A ``<!-- style: … -->`` directive on a list styles the *top level* through
   the style map (nested items keep ``List Number 2/3`` / ``List Bullet 2/3``)
   instead of flattening every produced paragraph to the directive style.

Assertions are made on the numbering/paragraph XML because Word computes the
visible numbers and indents at display time.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.markdown_processor import process_markdown_content  # noqa: E402
from docx_tools.numbering import resolve_ordered_numbering  # noqa: E402
from docx_tools.style_map import build_style_map  # noqa: E402


def _add_numbered_style(doc, name, num_fmt="upperRoman", lvl_left=2880,
                        lvl_hanging=720, style_left_inches=None):
    """Register paragraph style *name* backed by its own single-level numbering.

    Mimics a custom template's numbered style: a dedicated ``abstractNum`` (with
    its own numeral format and level indents), a ``num`` instance, and a style
    whose ``pPr`` references that instance. Returns the abstractNumId (str).
    """
    numbering = doc.part.numbering_part.element

    def _el(tag, **attrs):
        el = OxmlElement(tag)
        for key, value in attrs.items():
            el.set(qn('w:' + key), value)
        return el

    existing = [int(v) for v in numbering.xpath('./w:abstractNum/@w:abstractNumId')]
    abstract_id = max(existing) + 1 if existing else 0
    abstract = _el('w:abstractNum', abstractNumId=str(abstract_id))
    abstract.append(_el('w:multiLevelType', val='singleLevel'))
    lvl = _el('w:lvl', ilvl='0')
    lvl.append(_el('w:start', val='1'))
    lvl.append(_el('w:numFmt', val=num_fmt))
    lvl.append(_el('w:lvlText', val='%1.'))
    lvl.append(_el('w:lvlJc', val='left'))
    lvl_ppr = OxmlElement('w:pPr')
    lvl_ppr.append(_el('w:ind', left=str(lvl_left), hanging=str(lvl_hanging)))
    lvl.append(lvl_ppr)
    abstract.append(lvl)
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        numbering.append(abstract)
    num = numbering.add_num(abstract_id)

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style_ppr = style._element.get_or_add_pPr()
    num_pr = style_ppr.get_or_add_numPr()
    num_pr.get_or_add_numId().val = num.numId
    if style_left_inches is not None:
        style.paragraph_format.left_indent = Inches(style_left_inches)
    return str(abstract_id)


def _para(doc, text):
    for p in doc.paragraphs:
        if p.text == text:
            return p
    raise AssertionError(f"no paragraph with text {text!r}")


def _abstract_of_para(doc, paragraph):
    """Abstract numbering id behind the paragraph's direct ``numPr``, or None."""
    vals = paragraph._p.xpath('./w:pPr/w:numPr/w:numId/@w:val')
    if not vals:
        return None
    numbering = doc.part.numbering_part.element
    return str(numbering.num_having_numId(int(vals[0])).abstractNumId.val)


def _ilvl_of_para(paragraph):
    vals = paragraph._p.xpath('./w:pPr/w:numPr/w:ilvl/@w:val')
    return int(vals[0]) if vals else None


def _direct_ind(paragraph):
    """The paragraph's direct ``w:ind`` attributes (local names), or None."""
    inds = paragraph._p.xpath('./w:pPr/w:ind')
    if not inds:
        return None
    return {key.rsplit('}', 1)[-1]: value for key, value in inds[0].attrib.items()}


# ---------------------------------------------------------------------------
# Fix 1 — numbering resolves from the style actually applied
# ---------------------------------------------------------------------------

def test_mapped_list_number_style_uses_its_own_numbering():
    """A style_mapping'd ordered style restarts with ITS abstract, not List Number's."""
    doc = Document()
    abstract_id = _add_numbered_style(doc, "Legal Number")
    style_map = build_style_map({"list_number": "Legal Number"})

    process_markdown_content(doc, "1. one\n2. two\n", style_map=style_map)

    one, two = _para(doc, "one"), _para(doc, "two")
    assert one.style.name == "Legal Number"
    assert _abstract_of_para(doc, one) == abstract_id
    assert _abstract_of_para(doc, two) == abstract_id
    assert _ilvl_of_para(one) == 0


def test_directive_list_uses_styles_own_numbering():
    """A directive-styled numbered list restarts with the directive style's abstract."""
    doc = Document()
    abstract_id = _add_numbered_style(doc, "Legal Number")

    process_markdown_content(doc, "<!-- style: Legal Number -->\n1. one\n2. two\n")

    one = _para(doc, "one")
    assert one.style.name == "Legal Number"
    assert _abstract_of_para(doc, one) == abstract_id
    # The restart instance carries a startOverride at the style's declared level.
    num_id = one._p.xpath('./w:pPr/w:numPr/w:numId/@w:val')[0]
    numbering = doc.part.numbering_part.element
    num = numbering.num_having_numId(int(num_id))
    assert num.xpath('./w:lvlOverride[@w:ilvl="0"]/w:startOverride/@w:val') == ["1"]


def test_nested_level_uses_level_styles_own_numbering():
    """In a template with single-level abstracts, nested items must not point at
    an undefined ilvl of the parent's abstract (Word would render no number)."""
    doc = Document()  # blank python-docx template: every abstract is single-level

    process_markdown_content(doc, "1. parent\n   1. child\n")

    parent, child = _para(doc, "parent"), _para(doc, "child")
    assert child.style.name == "List Number 2"
    # The child resolves List Number 2's own numbering (declared ilvl 0 here),
    # on a different abstract than the parent's.
    numbering = doc.part.numbering_part.element
    style_el = doc.styles["List Number 2"]._element
    style_num = int(style_el.xpath('./w:pPr/w:numPr/w:numId/@w:val')[0])
    expected_abstract = str(numbering.num_having_numId(style_num).abstractNumId.val)
    assert _abstract_of_para(doc, child) == expected_abstract
    assert _abstract_of_para(doc, child) != _abstract_of_para(doc, parent)
    # The referenced ilvl must exist in the abstract definition.
    child_ilvl = _ilvl_of_para(child)
    defined = numbering.xpath(
        f'./w:abstractNum[@w:abstractNumId="{expected_abstract}"]/w:lvl/@w:ilvl')
    assert str(child_ilvl) in defined


def test_missing_custom_style_falls_back_to_builtin_numbering():
    """An unknown mapped style degrades to the List Number chain, not to nothing."""
    doc = Document()
    style_map = build_style_map({"list_number": "Nonexistent"})

    process_markdown_content(doc, "1. one\n2. two\n", style_map=style_map)

    one = _para(doc, "one")
    numbering = doc.part.numbering_part.element
    style_el = doc.styles["List Number"]._element
    style_num = int(style_el.xpath('./w:pPr/w:numPr/w:numId/@w:val')[0])
    expected_abstract = str(numbering.num_having_numId(style_num).abstractNumId.val)
    assert _abstract_of_para(doc, one) == expected_abstract


def test_resolver_prefers_style_name_over_builtins():
    doc = Document()
    abstract_id = _add_numbered_style(doc, "Legal Number")
    resolved, ilvl, _root = resolve_ordered_numbering(doc, level=0,
                                                      style_name="Legal Number")
    assert resolved == abstract_id
    assert ilvl == 0


def test_decimal_fallback_scan_survives_unregistered_abstract_elements():
    """Priority-3 fallback: no usable list styles, but a decimal abstractNum exists.

    ``w:abstractNum`` has no registered oxml class, so ``w:``-prefixed xpath on
    it raises XPathEvalError — the scan must use qualified-name lookups. Deleting
    the built-in numbered styles forces resolution onto this path.
    """
    doc = Document()
    doc.part.numbering_part  # materialise the numbering part (holds decimal abstracts)
    for name in ("List Number", "List Number 2", "List Number 3"):
        doc.styles[name].delete()

    abstract_id, ilvl, root = resolve_ordered_numbering(doc, level=0)

    assert abstract_id is not None and root is not None
    assert ilvl == 0
    # The returned abstract really is decimal at ilvl 0.
    fmts = root.xpath(
        f'./w:abstractNum[@w:abstractNumId="{abstract_id}"]'
        f'/w:lvl[@w:ilvl="0"]/w:numFmt/@w:val')
    assert fmts == ["decimal"]
    # And rendering still attaches restart numbering end-to-end.
    process_markdown_content(doc, "1. one\n2. two\n")
    assert _abstract_of_para(doc, _para(doc, "one")) == str(abstract_id)


def test_resolver_follows_based_on_chain():
    """A style without its own numPr inherits the numbering of its base style."""
    doc = Document()
    abstract_id = _add_numbered_style(doc, "Legal Number")
    derived = doc.styles.add_style("Legal Number Tight", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = doc.styles["Legal Number"]
    resolved, ilvl, _root = resolve_ordered_numbering(doc, level=0,
                                                      style_name="Legal Number Tight")
    assert resolved == abstract_id
    assert ilvl == 0


# ---------------------------------------------------------------------------
# Fix 2 — the style's indents survive the direct numPr
# ---------------------------------------------------------------------------

def test_style_indent_reasserted_on_ordered_items():
    """The style's w:ind is restated as direct formatting on every ordered item."""
    doc = Document()
    _add_numbered_style(doc, "Legal Number", style_left_inches=1.0)

    process_markdown_content(doc, "<!-- style: Legal Number -->\n1. one\n2. two\n")

    for text in ("one", "two"):
        ind = _direct_ind(_para(doc, text))
        assert ind is not None, f"item {text!r} lost the style's indent"
        assert ind.get("left") == str(int(Inches(1.0).twips))
        # The style defines no hanging — the numbering level's hanging still
        # applies (w:ind attributes merge individually), so none is stamped.
        assert "hanging" not in ind


def test_no_direct_indent_when_style_defines_none():
    """Without a styled indent, no direct w:ind is invented (template's numbering
    level keeps full control, as before)."""
    doc = Document()  # blank template: List Number style has no w:ind

    process_markdown_content(doc, "1. one\n2. two\n")

    assert _direct_ind(_para(doc, "one")) is None


def test_default_template_ordered_items_get_style_indent():
    """The shipped default template styles List Number at 576 twips; items must
    carry it directly so it beats the abstract's 432-twip level indent."""
    default = project_root / "default_templates" / "default_docx_template.docx"
    doc = Document(str(default))
    style_ind = doc.styles["List Number"]._element.xpath('./w:pPr/w:ind/@w:left')
    assert style_ind, "fixture expectation: default template styles List Number's indent"

    process_markdown_content(doc, "1. one\n")

    ind = _direct_ind(_para(doc, "one"))
    assert ind is not None and ind.get("left") == style_ind[0]


# ---------------------------------------------------------------------------
# Fix 3 — directive styles lists via the style map, keeping nested levels
# ---------------------------------------------------------------------------

def test_directive_list_keeps_nested_number_styles():
    doc = Document()
    _add_numbered_style(doc, "Legal Number")

    process_markdown_content(
        doc, "<!-- style: Legal Number -->\n1. parent\n   1. child\n")

    assert _para(doc, "parent").style.name == "Legal Number"
    assert _para(doc, "child").style.name == "List Number 2"


def test_directive_bullet_list_keeps_nested_bullet_styles():
    doc = Document()
    doc.styles.add_style("Fancy Bullet", WD_STYLE_TYPE.PARAGRAPH)

    process_markdown_content(
        doc, "<!-- style: Fancy Bullet -->\n- parent\n  - child\n")

    assert _para(doc, "parent").style.name == "Fancy Bullet"
    assert _para(doc, "child").style.name == "List Bullet 2"


def test_directive_on_nongenuine_numbered_line_still_styles_paragraph():
    """A numbered line that renders as prose (a standalone date) is still styled."""
    doc = Document()
    doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)

    process_markdown_content(doc, "<!-- style: Callout -->\n23. brezna 2026\n")

    para = _para(doc, "23. brezna 2026")
    assert para.style.name == "Callout"
    assert _ilvl_of_para(para) is None, "must not have been rendered as a list"


def test_directive_list_in_detached_mode_keeps_nested_styles():
    """The placeholder path (return_elements=True) gets the same level styling."""
    from docx.text.paragraph import Paragraph
    doc = Document()
    _add_numbered_style(doc, "Legal Number")

    elements = process_markdown_content(
        doc, "<!-- style: Legal Number -->\n1. parent\n   1. child\n",
        return_elements=True)

    styles = [Paragraph(e, doc._body).style.name for e in elements
              if e.tag == qn('w:p')]
    assert styles == ["Legal Number", "List Number 2"]
