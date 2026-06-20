"""Tests for conditional block resolution in dynamic DOCX templates.

Covers the Phase 1 block-level feature: keeping/dropping ranges of body content
between {{#if flag}} / {{/if}} markers (and the negated {{^if flag}} form).

Output files are saved to tests/output/docx/ for manual inspection.
"""

import sys
from pathlib import Path

# Add project root to path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

import pytest
from docx import Document

from docx_tools.conditionals import resolve_conditionals
from docx_tools.dynamic_docx_tools import _replace_placeholders_in_document

OUTPUT_DIR = Path(__file__).parent / "output" / "docx"


def setup_module(module):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def build_doc(lines):
    """Create a document with one paragraph per line."""
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    return doc


def body_texts(doc):
    """Non-empty paragraph texts in the document body, in order."""
    return [p.text for p in doc.paragraphs if p.text != ""]


# ---------------------------------------------------------------------------
# Basic keep / drop
# ---------------------------------------------------------------------------

def test_if_true_keeps_inner_and_strips_markers():
    doc = build_doc(["Before", "{{#if flag}}", "Inner", "{{/if}}", "After"])
    resolve_conditionals(doc, {"flag": True})
    assert body_texts(doc) == ["Before", "Inner", "After"]


def test_if_false_drops_inner_and_markers():
    doc = build_doc(["Before", "{{#if flag}}", "Inner", "{{/if}}", "After"])
    resolve_conditionals(doc, {"flag": False})
    assert body_texts(doc) == ["Before", "After"]


def test_if_false_drops_multiple_inner_paragraphs():
    doc = build_doc(
        ["Before", "{{#if flag}}", "One", "Two", "Three", "{{/if}}", "After"]
    )
    resolve_conditionals(doc, {"flag": False})
    assert body_texts(doc) == ["Before", "After"]


# ---------------------------------------------------------------------------
# Negation ({{^if}})
# ---------------------------------------------------------------------------

def test_unless_keeps_when_flag_false():
    doc = build_doc(["{{^if flag}}", "Shown when false", "{{/if}}"])
    resolve_conditionals(doc, {"flag": False})
    assert body_texts(doc) == ["Shown when false"]


def test_unless_drops_when_flag_true():
    doc = build_doc(["{{^if flag}}", "Shown when false", "{{/if}}"])
    resolve_conditionals(doc, {"flag": True})
    assert body_texts(doc) == []


# ---------------------------------------------------------------------------
# Nesting
# ---------------------------------------------------------------------------

def test_nested_outer_true_inner_false():
    doc = build_doc(
        [
            "{{#if outer}}",
            "outer-top",
            "{{#if inner}}",
            "inner",
            "{{/if}}",
            "outer-bottom",
            "{{/if}}",
        ]
    )
    resolve_conditionals(doc, {"outer": True, "inner": False})
    assert body_texts(doc) == ["outer-top", "outer-bottom"]


def test_sibling_blocks_evaluated_independently():
    # Two separate (non-nested) blocks with different names must be independent.
    doc = build_doc(
        ["{{#if a}}", "Section A", "{{/if}}", "{{#if b}}", "Section B", "{{/if}}"]
    )
    resolve_conditionals(doc, {"a": True, "b": False})
    assert body_texts(doc) == ["Section A"]


def test_nested_outer_false_drops_everything():
    doc = build_doc(
        [
            "{{#if outer}}",
            "outer-top",
            "{{#if inner}}",
            "inner",
            "{{/if}}",
            "{{/if}}",
        ]
    )
    resolve_conditionals(doc, {"outer": False, "inner": True})
    assert body_texts(doc) == []


# ---------------------------------------------------------------------------
# Whole tables between markers are pruned with the block range
# ---------------------------------------------------------------------------

def test_table_inside_false_block_is_dropped():
    doc = Document()
    doc.add_paragraph("Before")
    doc.add_paragraph("{{#if flag}}")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "in-table"
    doc.add_paragraph("{{/if}}")
    doc.add_paragraph("After")

    resolve_conditionals(doc, {"flag": False})

    assert body_texts(doc) == ["Before", "After"]
    assert len(doc.tables) == 0


def test_table_inside_true_block_is_kept():
    doc = Document()
    doc.add_paragraph("{{#if flag}}")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "in-table"
    doc.add_paragraph("{{/if}}")

    resolve_conditionals(doc, {"flag": True})

    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "in-table"


# ---------------------------------------------------------------------------
# Multi-run split markers (Word fragments typed text across runs)
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("opener", ["{{#if flag }}", "{{ #if flag}}", "{{ #if  flag }}"])
def test_whitespace_inside_braces_is_tolerated(opener):
    doc = build_doc([opener, "Inner", "{{/if}}"])
    resolve_conditionals(doc, {"flag": True})
    assert body_texts(doc) == ["Inner"]


def test_marker_split_across_runs_is_detected():
    doc = Document()
    doc.add_paragraph("Before")
    p = doc.add_paragraph()
    for chunk in ["{{#", "if ", "flag", "}}"]:
        p.add_run(chunk)
    doc.add_paragraph("Inner")
    doc.add_paragraph("{{/if}}")
    doc.add_paragraph("After")

    resolve_conditionals(doc, {"flag": False})

    assert body_texts(doc) == ["Before", "After"]


# ---------------------------------------------------------------------------
# Condition value handling
# ---------------------------------------------------------------------------

def test_missing_condition_in_if_keeps_content():
    doc = build_doc(["{{#if unknown}}", "Inner", "{{/if}}"])
    resolve_conditionals(doc, {})  # name not present -> keep
    assert body_texts(doc) == ["Inner"]


def test_missing_condition_in_unless_also_keeps_content():
    # Unknown name must keep content for {{^if}} too, not drop it.
    doc = build_doc(["{{^if unknown}}", "Inner", "{{/if}}"])
    resolve_conditionals(doc, {})
    assert body_texts(doc) == ["Inner"]


def test_none_value_is_falsy():
    doc = build_doc(["{{#if flag}}", "Inner", "{{/if}}"])
    resolve_conditionals(doc, {"flag": None})
    assert body_texts(doc) == []


# ---------------------------------------------------------------------------
# Forgiving error handling: unbalanced markers -> warn & keep content
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("flag", [True, False])
def test_unbalanced_open_keeps_content_regardless_of_flag(flag):
    doc = build_doc(["{{#if flag}}", "Inner", "still here"])  # no close
    resolve_conditionals(doc, {"flag": flag})
    # Content preserved either way; recognisable marker paragraph stripped.
    assert body_texts(doc) == ["Inner", "still here"]


def test_unbalanced_stray_close_keeps_content():
    doc = build_doc(["Inner", "{{/if}}", "After"])  # close without open
    resolve_conditionals(doc, {"flag": False})
    assert body_texts(doc) == ["Inner", "After"]


# ---------------------------------------------------------------------------
# No markers -> document untouched
# ---------------------------------------------------------------------------

def test_document_without_markers_unchanged():
    doc = build_doc(["One", "Two", "Three"])
    resolve_conditionals(doc, {"flag": True})
    assert body_texts(doc) == ["One", "Two", "Three"]


# ---------------------------------------------------------------------------
# Integration: conditional resolution + placeholder substitution (the order
# wired in dynamic_docx_tools._sync_impl)
# ---------------------------------------------------------------------------

def _render(doc, payload):
    """Mirror the _sync_impl pipeline: prune conditionals, then substitute."""
    resolve_conditionals(doc, payload)
    context = {k: ("" if v is None else str(v)) for k, v in payload.items()}
    _replace_placeholders_in_document(doc, context)


def test_pipeline_fills_placeholders_in_kept_block_only():
    doc = build_doc(
        [
            "{{#if greet}}",
            "Hello {{name}}!",
            "{{/if}}",
            "{{#if secret}}",
            "Secret for {{name}}",
            "{{/if}}",
            "Signed: {{name}}",
        ]
    )
    _render(doc, {"greet": True, "secret": False, "name": "Dan"})

    texts = body_texts(doc)
    assert "Hello Dan!" in texts          # kept block, placeholder filled
    assert "Signed: Dan" in texts         # outside any block, placeholder filled
    assert all("Secret" not in t for t in texts)   # dropped block, gone entirely
    assert all("{{" not in t for t in texts)       # no leftover markers/placeholders


# ---------------------------------------------------------------------------
# Visual-inspection document
# ---------------------------------------------------------------------------

def test_visual_inspection_document():
    doc = Document()
    doc.add_heading("Conditional Blocks Demo", level=1)
    doc.add_paragraph("This clause is always present.")
    doc.add_paragraph("{{#if include_arbitration}}")
    doc.add_paragraph("Arbitration: disputes resolved by binding arbitration.")
    doc.add_paragraph("{{/if}}")
    doc.add_paragraph("{{^if include_arbitration}}")
    doc.add_paragraph("Jurisdiction: disputes resolved in the local courts.")
    doc.add_paragraph("{{/if}}")
    doc.add_paragraph("Signature: ______________________")

    resolve_conditionals(doc, {"include_arbitration": True})

    texts = body_texts(doc)
    assert "Arbitration: disputes resolved by binding arbitration." in texts
    assert "Jurisdiction: disputes resolved in the local courts." not in texts

    output_path = OUTPUT_DIR / "conditional_blocks_demo.docx"
    doc.save(str(output_path))
    print(f"Saved: {output_path}")
