"""Regression tests for literal newline escape sequences and backslash escaping.

Models frequently emit a newline as the two literal characters ``\\n`` (or
``\\r\\n``) inside a tool argument instead of a real newline — usually because a
tool/argument description demonstrated ``\\n`` as if it were syntax. The old
backslash handler (``re.compile(r'\\(.)')``) stripped the slash before *any*
character, so a literal ``\\n`` collapsed to a stray ``n`` and the line break was
lost (it also corrupted ``\\t``, Windows paths like ``C:\\new``, etc.).

The fix:
  * ``_ESCAPE_RE`` now only escapes ASCII punctuation (CommonMark behaviour), so a
    backslash before a letter/digit is preserved instead of silently dropped.
  * ``normalize_escaped_newlines`` converts literal ``\\n``/``\\r\\n``/``\\r`` to
    real newlines so they render as genuine line/paragraph breaks.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.inline_formatting import parse_inline_formatting  # noqa: E402
from docx_tools.markdown_processor import process_markdown_content  # noqa: E402
from docx_tools.patterns import normalize_escaped_newlines  # noqa: E402
from docx_tools.dynamic_docx_tools import (  # noqa: E402
    _replace_placeholders_in_document,
)


def _render_inline(value):
    doc = Document()
    para = doc.add_paragraph()
    parse_inline_formatting(value, para)
    return para


def _break_count(paragraph):
    return paragraph._p.xml.count("<w:br")


def _default_doc():
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"
    return Document(str(default))


# --- normalize_escaped_newlines (unit) --------------------------------------

def test_normalize_literal_n():
    assert normalize_escaped_newlines(r"a\nb") == "a\nb"


def test_normalize_literal_crlf_is_single_break():
    assert normalize_escaped_newlines(r"a\r\nb") == "a\nb"


def test_normalize_literal_cr():
    assert normalize_escaped_newlines(r"a\rb") == "a\nb"


def test_normalize_leaves_real_newline_untouched():
    assert normalize_escaped_newlines("a\nb") == "a\nb"


def test_normalize_does_not_touch_other_backslash_letters():
    # \t / \d are not newline escapes — they must be left for the (punctuation-
    # only) escape handler, which now preserves the backslash.
    assert normalize_escaped_newlines(r"a\tb") == r"a\tb"


# --- the reported failure: literal \n collapsing to "n" ---------------------

def test_literal_newline_does_not_become_stray_n():
    para = _render_inline(r'(dále jen „Dlužník").\n\n\nMezi Klientem')
    assert "nnn" not in para.text
    assert para.text == '(dále jen „Dlužník").\n\n\nMezi Klientem'


def test_literal_newline_renders_as_break():
    para = _render_inline(r"Line one\nLine two")
    assert _break_count(para) == 1
    assert para.text == "Line one\nLine two"


def test_real_newline_still_renders_as_break():
    para = _render_inline("Line one\nLine two")
    assert _break_count(para) == 1


# --- backslash escaping must still work for markdown punctuation -------------

def test_escaped_dot_still_works():
    assert _render_inline(r"1\. ledna 2026").text == "1. ledna 2026"


def test_escaped_asterisk_still_works():
    assert _render_inline(r"\*not italic\*").text == "*not italic*"


def test_escaped_backslash_still_works():
    assert _render_inline(r"a\\b").text == "a\\b"


def test_escaped_backtick_still_works():
    assert _render_inline(r"\`code\`").text == "`code`"


# --- no more silent corruption of backslash-before-letter -------------------

def test_backslash_before_letter_is_preserved():
    # Previously "a\tb" -> "atb" (backslash and the escape meaning lost). The
    # backslash must now survive since \t is not a markdown escape.
    assert _render_inline(r"a\tb").text == r"a\tb"


# --- block path -------------------------------------------------------------

def test_block_literal_double_newline_makes_separate_paragraphs():
    doc = Document()
    process_markdown_content(doc, r"First para.\n\nSecond para.")
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts == ["First para.", "Second para."]


def test_block_literal_newline_list_is_detected():
    doc = _default_doc()
    start = len(doc.paragraphs)
    process_markdown_content(doc, r"1. First item\n2. Second item")
    new = [p for p in doc.paragraphs[start:] if p.text.strip()]
    assert [p.text for p in new] == ["First item", "Second item"]
    assert all(p.style.name.startswith("List Number") for p in new)


# --- template placeholder path (the actual failing surface) -----------------

def test_template_placeholder_literal_newline():
    doc = Document()
    doc.add_paragraph().add_run("{{body}}")
    _replace_placeholders_in_document(
        doc, {"body": r"First paragraph.\n\nSecond paragraph."}
    )
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "nn" not in full  # no stray collapsed-newline characters
    assert "First paragraph." in full
    assert "Second paragraph." in full


# --- template line-break model: blank line = paragraph, <br> = soft break -----

def _fill_template(value):
    doc = Document()
    doc.add_paragraph().add_run("{{body}}")
    _replace_placeholders_in_document(doc, {"body": value})
    return doc


def test_template_blank_line_makes_separate_paragraphs():
    doc = _fill_template("First paragraph.\n\nSecond paragraph.")
    bodies = [p for p in doc.paragraphs if p.text.strip()]
    assert [p.text for p in bodies] == ["First paragraph.", "Second paragraph."]


def test_template_literal_blank_line_makes_separate_paragraphs():
    # The real failing shape: blank line written as literal "\n\n".
    doc = _fill_template(r"First paragraph.\n\nSecond paragraph.")
    bodies = [p for p in doc.paragraphs if p.text.strip()]
    assert [p.text for p in bodies] == ["First paragraph.", "Second paragraph."]


def test_template_br_is_soft_break_single_paragraph():
    doc = _fill_template("Line one<br>Line two")
    bodies = [p for p in doc.paragraphs if p.text.strip()]
    assert len(bodies) == 1
    assert bodies[0]._p.xml.count("<w:br") == 1


def test_template_multi_paragraph_inherits_placeholder_run_format():
    # Bold placeholder run -> every produced paragraph should stay bold.
    doc = Document()
    run = doc.add_paragraph().add_run("{{body}}")
    run.bold = True
    _replace_placeholders_in_document(doc, {"body": "First.\n\nSecond."})
    bodies = [p for p in doc.paragraphs if p.text.strip()]
    assert len(bodies) == 2
    assert all(r.bold for p in bodies for r in p.runs)


# --- unified newline semantics: template == base tool -----------------------

def test_template_single_newline_is_new_paragraph():
    # A lone newline now starts a new paragraph in templates too (was a soft
    # break before), matching the base tool.
    doc = Document()
    doc.add_paragraph().add_run("{{x}}")
    _replace_placeholders_in_document(doc, {"x": "A\nB"})
    bodies = [p for p in doc.paragraphs if p.text.strip()]
    assert [p.text for p in bodies] == ["A", "B"]


def test_template_matches_base_pipeline_for_multiline_prose():
    md = "First line.\nSecond line.\n\nThird paragraph."
    base_doc = Document()
    process_markdown_content(base_doc, md)
    base_sig = [(p.style.name, p.text) for p in base_doc.paragraphs if p.text.strip()]

    tmpl_doc = Document()
    tmpl_doc.add_paragraph().add_run("{{x}}")
    _replace_placeholders_in_document(tmpl_doc, {"x": md})
    tmpl_sig = [(p.style.name, p.text) for p in tmpl_doc.paragraphs if p.text.strip()]

    assert tmpl_sig == base_sig
    assert len(tmpl_sig) == 3


def test_template_single_line_value_stays_in_place():
    # A single-line value must NOT spawn a new paragraph (kept inline so a
    # mid-paragraph "Dear {{name}}," works and the original paragraph is reused).
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("{{x}}")
    _replace_placeholders_in_document(doc, {"x": "just one line"})
    bodies = [p for p in doc.paragraphs if p.text.strip()]
    assert len(bodies) == 1
    assert bodies[0]._p is para._p  # same underlying paragraph element reused
    assert bodies[0].text == "just one line"


def test_template_prose_inherits_alignment_but_heading_keeps_style():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("{{x}}")
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _replace_placeholders_in_document(
        doc, {"x": "Odstavec jedna.\n\n## Nadpis\n\nOdstavec dva."}
    )
    by_text = {p.text: p for p in doc.paragraphs if p.text.strip()}
    # plain prose inherits the placeholder paragraph's justification
    assert by_text["Odstavec jedna."].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert by_text["Odstavec dva."].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    # the heading keeps its own style and is NOT forced to justify
    assert by_text["Nadpis"].style.name == "Heading 2"
    assert by_text["Nadpis"].alignment != WD_ALIGN_PARAGRAPH.JUSTIFY
