"""Tests for flexible list-nesting indentation (improvement 4.2).

Nesting is determined by *relative* indentation, so 2-, 3- and 4-space units
(and tabs) all nest correctly — previously only a 3-space step worked.
"""
import sys
from pathlib import Path

from docx import Document

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.markdown_processor import process_markdown_content  # noqa: E402


def _doc():
    default = project_root / "default_templates" / "default_docx_template.docx"
    return Document(str(default))


def _styles(doc, start):
    return [(p.text, p.style.name) for p in doc.paragraphs[start:] if p.text]


import pytest  # noqa: E402


@pytest.mark.parametrize("unit", ["  ", "   ", "    ", "\t"])
def test_two_three_four_space_and_tab_all_nest(unit):
    doc = _doc()
    start = len(doc.paragraphs)
    md = f"- parent\n{unit}- child\n{unit}{unit}- grandchild\n"
    process_markdown_content(doc, md)
    assert _styles(doc, start) == [
        ("parent", "List Bullet"),
        ("child", "List Bullet 2"),
        ("grandchild", "List Bullet 3"),
    ]


def test_ordered_two_space_nesting():
    doc = _doc()
    start = len(doc.paragraphs)
    md = "1. a\n  1. a1\n  2. a2\n2. b\n"
    process_markdown_content(doc, md)
    assert _styles(doc, start) == [
        ("a", "List Number"),
        ("a1", "List Number 2"),
        ("a2", "List Number 2"),
        ("b", "List Number"),
    ]


def test_dedent_returns_to_parent_level():
    doc = _doc()
    start = len(doc.paragraphs)
    md = "- a\n  - a1\n- b\n"  # back out to top level
    process_markdown_content(doc, md)
    assert _styles(doc, start) == [
        ("a", "List Bullet"),
        ("a1", "List Bullet 2"),
        ("b", "List Bullet"),
    ]


def test_overindented_child_is_one_level_deeper():
    # A child indented far more than one step is still just one level deeper,
    # matching how real markdown parsers treat over-indentation.
    doc = _doc()
    start = len(doc.paragraphs)
    md = "- a\n      - deep\n"  # 6 spaces
    process_markdown_content(doc, md)
    assert _styles(doc, start) == [
        ("a", "List Bullet"),
        ("deep", "List Bullet 2"),
    ]
