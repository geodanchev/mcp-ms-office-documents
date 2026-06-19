"""Tests for the ad-hoc style directive (issue #66, part B).

``<!-- style: Name -->`` applies the named Word style to the next rendered block
(paragraph, heading, list — every item — or table). It overrides the active
StyleMap for that block only, and falls back gracefully when the style is missing.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.markdown_processor import process_markdown_content  # noqa: E402


def _doc(*para_styles, table_styles=()):
    """Blank doc with the given custom paragraph/table styles registered."""
    d = Document()
    for name in para_styles:
        d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    for name in table_styles:
        d.styles.add_style(name, WD_STYLE_TYPE.TABLE)
    return d


def _style_of(doc, start, text):
    for p in doc.paragraphs[start:]:
        if p.text == text:
            return p.style.name
    raise AssertionError(f"no paragraph with text {text!r}")


def test_directive_styles_following_paragraph():
    doc = _doc("Callout")
    start = len(doc.paragraphs)
    process_markdown_content(doc, "<!-- style: Callout -->\nhello world\n")
    assert _style_of(doc, start, "hello world") == "Callout"


def test_directive_styles_every_list_item():
    doc = _doc("Fancy List")
    start = len(doc.paragraphs)
    process_markdown_content(doc, "<!-- style: Fancy List -->\n1. one\n2. two\n3. three\n")
    for text in ("one", "two", "three"):
        assert _style_of(doc, start, text) == "Fancy List"


def test_directive_only_affects_the_next_block():
    doc = _doc("Callout")
    start = len(doc.paragraphs)
    process_markdown_content(doc, "<!-- style: Callout -->\nstyled\n\nplain\n")
    assert _style_of(doc, start, "styled") == "Callout"
    assert _style_of(doc, start, "plain") == "Normal"


def test_directive_applies_to_table():
    doc = _doc(table_styles=("Fancy Grid",))
    process_markdown_content(doc, "<!-- style: Fancy Grid -->\n| A | B |\n|---|---|\n| 1 | 2 |\n")
    assert doc.tables[-1].style.name == "Fancy Grid"


def test_directive_blank_line_between_directive_and_block():
    doc = _doc("Callout")
    start = len(doc.paragraphs)
    process_markdown_content(doc, "<!-- style: Callout -->\n\nhello\n")
    assert _style_of(doc, start, "hello") == "Callout"


def test_directive_is_case_insensitive():
    doc = _doc("Callout")
    start = len(doc.paragraphs)
    process_markdown_content(doc, "<!-- STYLE: Callout -->\nhello\n")
    assert _style_of(doc, start, "hello") == "Callout"


def test_unknown_style_falls_back_without_failing():
    doc = Document()
    start = len(doc.paragraphs)
    process_markdown_content(doc, "<!-- style: Nonexistent -->\ntext\n")
    assert _style_of(doc, start, "text") == "Normal"  # fell back, no exception


def test_trailing_directive_is_a_noop():
    doc = Document()
    # Directive with nothing after it must not raise or hang.
    process_markdown_content(doc, "intro\n\n<!-- style: Callout -->\n")
    assert doc.paragraphs[-1].text == "intro"


def test_directive_in_detached_mode():
    doc = _doc("Callout")
    elements = process_markdown_content(
        doc, "<!-- style: Callout -->\n1. a\n2. b\n", return_elements=True
    )
    styled = [Paragraph(e, doc._body) for e in elements if e.tag.endswith('}p')]
    assert styled and all(p.style.name == "Callout" for p in styled)


def test_stacked_directives_compose_on_a_table():
    """style + borderless + widths can all be stacked above one table (4.6)."""
    doc = _doc(table_styles=("Fancy Grid",))
    md = (
        "<!-- style: Fancy Grid -->\n"
        "<!-- borderless -->\n"
        "<!-- widths: 30 70 -->\n"
        "| A | B |\n|---|---|\n| 1 | 2 |\n"
    )
    process_markdown_content(doc, md)
    table = doc.tables[-1]
    assert table.style.name == "Fancy Grid"
    # 30/70 split → second column noticeably wider than the first.
    assert table.rows[0].cells[1].width > table.rows[0].cells[0].width


def test_nondirective_comment_is_skipped_not_rendered():
    doc = Document()
    start = len(doc.paragraphs)
    process_markdown_content(doc, "before\n\n<!-- just a free-text note -->\n\nafter\n")
    texts = [p.text for p in doc.paragraphs[start:] if p.text]
    assert texts == ["before", "after"]  # the comment is not rendered as a paragraph
