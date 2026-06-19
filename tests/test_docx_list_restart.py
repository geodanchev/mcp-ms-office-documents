"""Tests for numbered-list restart (issue #67).

When a numbered list restarts — a new list, or ``1.`` reappearing at the same
level — Word should count from 1 again instead of continuing the previous list.
The renderer implements this by giving each logical list its own ``<w:num>``
instance with a ``<w:startOverride>`` (see ``docx_tools/numbering.py``), rather
than sharing the ``List Number`` style's single numbering definition.

These tests assert on the underlying numbering XML because the visible numbers
are computed by Word at display time and are not stored as paragraph text.
"""
import sys
from pathlib import Path

from docx import Document

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.block_elements import process_list_items  # noqa: E402
from docx_tools.numbering import resolve_ordered_abstract_num_id  # noqa: E402


def _new_doc_with_default_styles():
    """Document from the project's default template (carries ``List Number``)."""
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"
    return Document(str(default))


def _num_id_of(paragraph):
    """Return the explicit numId attached to *paragraph*, or ``None``."""
    vals = paragraph._p.xpath('.//w:numPr/w:numId/@w:val')
    return vals[0] if vals else None


def _start_override(doc, num_id):
    """Return the ilvl-0 startOverride value for *num_id* in the numbering part."""
    numbering = doc.part.numbering_part.element
    num = numbering.num_having_numId(int(num_id))
    vals = num.xpath('./w:lvlOverride[@w:ilvl="0"]/w:startOverride/@w:val')
    return vals[0] if vals else None


def _ordered_paragraphs(doc, start):
    return [
        p for p in doc.paragraphs[start:]
        if p.style.name and p.style.name.startswith("List Number")
    ]


def test_two_adjacent_lists_get_distinct_numbering_instances():
    """A second list starting at ``1.`` restarts instead of continuing."""
    lines = [
        "1. Apple",
        "2. Banana",
        "",
        "1. Carrot",
        "2. Potato",
    ]
    doc = _new_doc_with_default_styles()
    start = len(doc.paragraphs)

    process_list_items(lines, 0, doc, is_ordered=True, level=0)

    paras = _ordered_paragraphs(doc, start)
    assert [p.text for p in paras] == ["Apple", "Banana", "Carrot", "Potato"]

    first_num = _num_id_of(paras[0])
    second_num = _num_id_of(paras[2])
    assert first_num is not None and second_num is not None, (
        "ordered items should carry an explicit numId"
    )
    # Same list shares a numId; the restart gets a fresh one.
    assert _num_id_of(paras[1]) == first_num
    assert _num_id_of(paras[3]) == second_num
    assert first_num != second_num, "second list must restart on a new numId"
    assert _start_override(doc, second_num) == "1"


def test_list_starting_at_n_overrides_start_to_n():
    """A list that begins at ``5.`` starts the override at 5, not 1."""
    lines = ["5. Fifth", "6. Sixth"]
    doc = _new_doc_with_default_styles()
    start = len(doc.paragraphs)

    process_list_items(lines, 0, doc, is_ordered=True, level=0)

    paras = _ordered_paragraphs(doc, start)
    num_id = _num_id_of(paras[0])
    assert _start_override(doc, num_id) == "5"


def test_nested_ordered_list_restarts_independently_of_parent():
    """A nested numbered list gets its own numId so it doesn't share the parent count."""
    lines = [
        "1. Parent one",
        "   1. Child one",
        "   2. Child two",
        "2. Parent two",
    ]
    doc = _new_doc_with_default_styles()
    start = len(doc.paragraphs)

    process_list_items(lines, 0, doc, is_ordered=True, level=0)

    paras = _ordered_paragraphs(doc, start)
    by_text = {p.text: p for p in paras}
    parent_num = _num_id_of(by_text["Parent one"])
    child_num = _num_id_of(by_text["Child one"])

    assert parent_num is not None and child_num is not None
    assert parent_num != child_num, "child list must use a distinct numId"
    assert _num_id_of(by_text["Parent two"]) == parent_num, (
        "parent items keep the same numId across the nested child"
    )
    assert _num_id_of(by_text["Child two"]) == child_num


def test_restart_survives_return_elements_roundtrip():
    """The placeholder path removes paragraphs from the body; numbering must persist."""
    lines = ["1. One", "2. Two", "", "1. Reset"]
    doc = _new_doc_with_default_styles()

    _idx, elements = process_list_items(
        lines, 0, doc, is_ordered=True, level=0, return_elements=True
    )

    assert elements, "return_elements should yield the detached paragraph elements"
    num_ids = [e.xpath('.//w:numPr/w:numId/@w:val') for e in elements]
    num_ids = [v[0] for v in num_ids if v]
    assert len(num_ids) == 3
    assert num_ids[0] == num_ids[1], "first two items share a numId"
    assert num_ids[2] != num_ids[0], "the reset item restarts on a fresh numId"
    # The <w:num> definitions live in the numbering part and outlive detachment.
    assert _start_override(doc, num_ids[2]) == "1"


def test_resolve_uses_list_number_style_abstract():
    """Restart instances reuse the template's ``List Number`` numbering format."""
    doc = _new_doc_with_default_styles()
    abstract_id, numbering = resolve_ordered_abstract_num_id(doc)
    assert abstract_id is not None
    # It should match what the List Number style points at.
    style_el = doc.styles["List Number"]._element
    style_num = style_el.xpath('.//w:numPr/w:numId/@w:val')[0]
    expected = numbering.num_having_numId(int(style_num)).abstractNumId.val
    assert str(abstract_id) == str(expected)


def test_resolve_returns_str_abstract_id():
    """All resolver paths return a str abstract id (consistent type)."""
    doc = _new_doc_with_default_styles()
    abstract_id, _ = resolve_ordered_abstract_num_id(doc)
    assert isinstance(abstract_id, str)


def test_ordered_capture_pattern_agrees_with_detection_on_empty_item():
    """ORDERED_LIST_CAPTURE_PATTERN matches whenever ORDERED_LIST_PATTERN does."""
    from docx_tools.patterns import ORDERED_LIST_PATTERN, ORDERED_LIST_CAPTURE_PATTERN
    s = "1. "  # marker with no item text
    assert ORDERED_LIST_PATTERN.match(s)
    m = ORDERED_LIST_CAPTURE_PATTERN.match(s)
    assert m and m.group(1) == "1" and m.group(2) == ""
