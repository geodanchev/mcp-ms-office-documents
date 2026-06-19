"""Tests for custom style mapping (issue #66, part A).

Covers the :class:`StyleMap` abstraction, config merging precedence, the
``apply_style`` fallback, and end-to-end application through
``process_markdown_content``.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.style_map import (  # noqa: E402
    DEFAULT_STYLE_MAP,
    StyleMap,
    build_style_map,
    apply_style,
)
from docx_tools.markdown_processor import process_markdown_content  # noqa: E402


def _doc_with_default_styles():
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"
    return Document(str(default))


def _styled(doc, start, text):
    for p in doc.paragraphs[start:]:
        if p.text == text:
            return p.style.name
    raise AssertionError(f"no paragraph with text {text!r}")


# --------------------------------------------------------------------------- #
# StyleMap / build_style_map
# --------------------------------------------------------------------------- #
def test_default_map_matches_builtin_names():
    assert DEFAULT_STYLE_MAP.heading_style(1) == "Heading 1"
    assert DEFAULT_STYLE_MAP.heading_style(6) == "Heading 6"
    assert DEFAULT_STYLE_MAP.list_number[0] == "List Number"
    assert DEFAULT_STYLE_MAP.list_bullet[1] == "List Bullet 2"
    assert DEFAULT_STYLE_MAP.quote == "Quote"
    assert DEFAULT_STYLE_MAP.table == "Table Grid"


def test_heading_level_clamps():
    assert DEFAULT_STYLE_MAP.heading_style(0) == "Heading 1"
    assert DEFAULT_STYLE_MAP.heading_style(9) == "Heading 6"


def test_build_style_map_empty_returns_default_singleton():
    assert build_style_map() is DEFAULT_STYLE_MAP
    assert build_style_map({}, None) is DEFAULT_STYLE_MAP


def test_build_style_map_applies_scalar_and_indexed_overrides():
    sm = build_style_map({
        "quote": "My Quote",
        "table": "My Grid",
        "heading_2": "My H2",
        "list_number": "My Numbers",
        "list_bullet_3": "Deep Bullet",
    })
    assert sm.quote == "My Quote"
    assert sm.table == "My Grid"
    assert sm.heading_style(2) == "My H2"
    assert sm.heading_style(1) == "Heading 1"  # untouched
    assert sm.list_number == ("My Numbers", "List Number 2", "List Number 3")
    assert sm.list_bullet[2] == "Deep Bullet"


def test_per_template_mapping_overrides_global():
    global_cfg = {"quote": "Global Quote", "heading_1": "Global H1"}
    template_cfg = {"quote": "Template Quote", "heading_2": "Template H2"}
    sm = build_style_map(global_cfg, template_cfg)
    # template wins on conflict…
    assert sm.quote == "Template Quote"
    # …and non-conflicting global + template overrides both survive the merge.
    assert sm.heading_style(1) == "Global H1"
    assert sm.heading_style(2) == "Template H2"


def test_unknown_keys_are_ignored():
    sm = build_style_map({"bogus_key": "x", "quote": "Kept"})
    assert sm.quote == "Kept"
    assert sm == build_style_map({"quote": "Kept"})


# --------------------------------------------------------------------------- #
# apply_style fallback
# --------------------------------------------------------------------------- #
def test_apply_style_falls_back_on_missing_style():
    doc = _doc_with_default_styles()
    para = doc.add_paragraph("x")
    apply_style(para, "Definitely Not A Style", fallback="Normal")
    assert para.style.name == "Normal"  # no exception, fell back


def test_apply_style_noop_on_empty_name():
    doc = _doc_with_default_styles()
    para = doc.add_paragraph("x")
    before = para.style.name
    apply_style(para, None)
    assert para.style.name == before


# --------------------------------------------------------------------------- #
# End-to-end through process_markdown_content
# --------------------------------------------------------------------------- #
def test_custom_styles_applied_to_rendered_markdown():
    doc = _doc_with_default_styles()
    doc.styles.add_style("Brand Numbers", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("Brand Quote", WD_STYLE_TYPE.PARAGRAPH)
    sm = build_style_map({"list_number": "Brand Numbers", "quote": "Brand Quote"})

    start = len(doc.paragraphs)
    process_markdown_content(doc, "1. one\n2. two\n\n> a quote\n", style_map=sm)

    assert _styled(doc, start, "one") == "Brand Numbers"
    assert _styled(doc, start, "two") == "Brand Numbers"
    assert _styled(doc, start, "a quote") == "Brand Quote"


def test_default_map_preserves_builtin_rendering():
    doc = _doc_with_default_styles()
    start = len(doc.paragraphs)
    process_markdown_content(doc, "# Title\n\n1. item\n\n- bullet\n\n> quote\n")

    assert _styled(doc, start, "Title") == "Heading 1"
    assert _styled(doc, start, "item") == "List Number"
    assert _styled(doc, start, "bullet") == "List Bullet"
    assert _styled(doc, start, "quote") == "Quote"


def test_missing_mapped_style_falls_back_without_failing():
    doc = _doc_with_default_styles()
    sm = StyleMap(quote="No Such Style")
    start = len(doc.paragraphs)
    process_markdown_content(doc, "> resilient quote\n", style_map=sm)
    # Rendered (didn't raise) and fell back to the document default.
    assert _styled(doc, start, "resilient quote") == "Normal"
